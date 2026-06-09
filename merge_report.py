import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Load old file as base
doc = Document('Documentations/BaoCao_DoAnTotNghiep_ChatbotAI.docx')

def find_paragraph_index(doc, search_text):
    for i, p in enumerate(doc.paragraphs):
        if search_text in p.text:
            return i
    return -1

def add_paragraph_after(doc, ref_index, text, style='Normal'):
    """Add paragraph after a specific index"""
    # We'll append at end and note position
    p = doc.add_paragraph(text, style=style)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    return table

# ── Find insertion points ──
idx_chuong2 = find_paragraph_index(doc, 'CHƯƠNG 2. PHÂN TÍCH')
idx_chuong3 = find_paragraph_index(doc, 'CHƯƠNG 3. THIẾT KẾ')
idx_chuong4 = find_paragraph_index(doc, 'CHƯƠNG 4. HIỆN THỰC')
idx_huong_phat_trien = find_paragraph_index(doc, '4.6. Hướng phát triển')
idx_tai_lieu = find_paragraph_index(doc, 'TÀI LIỆU THAM KHẢO')

print(f"Chuong 2: {idx_chuong2}")
print(f"Chuong 3: {idx_chuong3}")
print(f"Chuong 4: {idx_chuong4}")
print(f"Huong PT: {idx_huong_phat_trien}")
print(f"Tai lieu: {idx_tai_lieu}")

# ── Add new sections BEFORE Chapter 2 (as Chapter 2 new content) ──
# We'll add new content at the END of the document in organized sections

doc.add_page_break()
doc.add_heading('PHẦN BỔ SUNG: TÍNH NĂNG MỚI', level=1)

doc.add_paragraph(
    'Phần này bổ sung các tính năng mới được phát triển sau phiên bản báo cáo ban đầu, '
    'bao gồm: Hệ thống Camera RTSP, Quản lý sân, và Framework theory.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PHẦN 1: FRAMEWORK THEORY (bổ sung vào Chương 2)
# ══════════════════════════════════════════════════════════
doc.add_heading('BỔ SUNG 1: CƠ SỞ LÝ THUYẾT FRAMEWORK', level=1)
doc.add_paragraph('(Bổ sung vào Chương 2 - Cơ sở lý thuyết)')

doc.add_heading('1.1. Tổng quan Framework', level=2)
doc.add_paragraph(
    'Framework là bộ khung phần mềm cung cấp các thành phần, công cụ và quy tắc tái sử dụng '
    'để phát triển ứng dụng nhanh hơn, nhất quán hơn và dễ bảo trì hơn. Việc lựa chọn framework '
    'phù hợp giúp giảm thời gian phát triển, tăng tính ổn định và dễ dàng mở rộng hệ thống.'
)
doc.add_paragraph(
    'Hệ thống sử dụng 2 framework chính: Flutter (frontend) và FastAPI (backend), '
    'kết hợp với các thư viện (packages) chuyên biệt cho từng nghiệp vụ.'
)

doc.add_heading('1.2. Flutter Framework & Packages', level=2)

doc.add_heading('1.2.1. Flutter là gì?', level=3)
doc.add_paragraph(
    'Flutter là framework phát triển ứng dụng cross-platform do Google phát triển, '
    'sử dụng ngôn ngữ lập trình Dart. Flutter cho phép xây dựng ứng dụng cho iOS, Android, '
    'Web và Desktop từ một codebase duy nhất.'
)
doc.add_paragraph('Đặc điểm nổi bật của Flutter:')
for f in [
    'Hot Reload: Thay đổi code hiển thị ngay lập tức, tăng tốc development.',
    'Widget-based: Giao diện được xây dựng từ các widget có thể tái sử dụng.',
    'Performance: Render engine riêng (Skia/Impeller), không dùng native UI components.',
    'Cross-platform: Một codebase chạy được trên nhiều nền tảng.',
    'Rich ecosystem: Hàng nghìn packages trên pub.dev.',
]:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('1.2.2. So sánh State Management: Riverpod vs Provider vs Bloc', level=3)
doc.add_paragraph(
    'State management là bài toán quan trọng trong Flutter. Hệ thống sử dụng Riverpod 2.4.9.'
)
add_table(doc, ['Tiêu chí', 'Provider', 'Riverpod', 'Bloc'], [
    ['Compile-safe', '❌ Runtime error', '✅ Compile-time check', '✅ Compile-time check'],
    ['Testability', 'Khó (cần BuildContext)', 'Dễ (independent)', 'Dễ (stream-based)'],
    ['Boilerplate', 'Ít', 'Ít', 'Nhiều (events/states)'],
    ['Learning curve', 'Dễ', 'Trung bình', 'Khó'],
    ['Dependency injection', 'Cần context', 'Global access', 'Manual setup'],
    ['Auto dispose', 'Không', 'Có (autoDispose)', 'Manual'],
    ['Family provider', 'Không', 'Có (family)', 'Không'],
])
doc.add_paragraph(
    'Lựa chọn Riverpod vì: type-safe, không cần BuildContext để đọc state, '
    'hỗ trợ autoDispose và family providers, dễ test.'
)

doc.add_heading('1.2.3. Các Packages tiêu biểu', level=3)
add_table(doc, ['Package', 'Vai trò'], [
    ['flutter_riverpod', 'State management declarative'],
    ['go_router', 'Declarative navigation với role-based redirect'],
    ['dio', 'HTTP client với interceptors'],
    ['web_socket_channel', 'WebSocket real-time communication'],
    ['flutter_chat_ui', 'Chat UI components'],
    ['flutter_stripe', 'Stripe native payment integration'],
    ['media_kit + media_kit_video', 'RTSP/HLS video playback (libmpv)'],
    ['qr_flutter + mobile_scanner', 'QR code generation & scanning'],
    ['flutter_secure_storage', 'Encrypted storage cho tokens'],
    ['local_auth', 'Biometric/PIN authentication'],
    ['flutter_webrtc', 'Voice call giữa staff và customer'],
])

doc.add_heading('1.3. FastAPI Backend Framework', level=2)
doc.add_paragraph(
    'FastAPI là framework xây dựng API hiện đại bằng Python, '
    'tốc độ ngang với Node.js và Go nhờ sử dụng ASGI.'
)
doc.add_paragraph('Đặc điểm nổi bật:')
for f in [
    'Tốc độ cao: Sử dụng Starlette và Pydantic, async/await native.',
    'Type-safe: Tự động validate request/response dựa trên type hints.',
    'Auto documentation: Tự sinh Swagger UI và ReDoc từ code.',
    'Dependency Injection: Hệ thống dependency injection mạnh mẽ.',
    'Async support: Hỗ trợ asyncio cho I/O-bound operations.',
]:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('1.4. Hệ quản trị cơ sở dữ liệu', level=2)
doc.add_paragraph('Hệ thống sử dụng 3 loại cơ sở dữ liệu:')

doc.add_heading('1.4.1. PostgreSQL (Relational Database)', level=3)
doc.add_paragraph(
    'PostgreSQL là RDBMS mã nguồn mở, hỗ trợ ACID transactions, JSONB, full-text search. '
    'Vai trò: Lưu trữ dữ liệu chính (users, bookings, orders, venues, cameras).'
)

doc.add_heading('1.4.2. Neo4j (Graph Database)', level=3)
doc.add_paragraph(
    'Neo4j là Graph DBMS, lưu trữ dữ liệu dưới dạng nodes và relationships. '
    'Vai trò: Knowledge Graph tri thức thể thao, Graph RAG cho AI chatbot.'
)

doc.add_heading('1.4.3. Redis (In-memory Database)', level=3)
doc.add_paragraph(
    'Redis là in-memory database, hỗ trợ pub/sub, session storage. '
    'Vai trò: Cache, staff notifications, ephemeral chat rooms.'
)

doc.add_heading('1.4.4. So sánh 3 loại Database', level=3)
add_table(doc, ['Tiêu chí', 'PostgreSQL', 'Neo4j', 'Redis'], [
    ['Loại', 'Relational', 'Graph', 'In-memory'],
    ['Storage', 'Disk', 'Disk', 'RAM'],
    ['Speed', 'Nhanh (indexed)', 'Nhanh (graph traversal)', 'Rất nhanh'],
    ['ACID', '✅', '✅', '✅'],
    ['Use case', 'OLTP, CRUD', 'Knowledge graph', 'Cache, pub/sub'],
    ['Ngôn ngữ', 'SQL', 'Cypher', 'Redis commands'],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PHẦN 2: CAMERA SYSTEM
# ══════════════════════════════════════════════════════════
doc.add_heading('BỔ SUNG 2: HỆ THỐNG CAMERA RTSP', level=1)
doc.add_paragraph('(Bổ sung vào Chương 3 & 4 - Thiết kế và Triển khai)')

doc.add_heading('2.1. Tổng quan hệ thống Camera', level=2)
doc.add_paragraph(
    'Hệ thống camera RTSP cho phép nhân viên giám sát sân/bàn trực tiếp từ ứng dụng di động, '
    'giải quyết vấn đề khách bỏ đi mà không thanh toán.'
)

doc.add_heading('2.2. Kiến trúc Camera Streaming', level=2)
doc.add_paragraph('[IP Camera] --RTSP TCP--> [media_kit/libmpv] --> [Flutter Video Widget]')
doc.add_paragraph('Hỗ trợ các hãng camera:')
add_table(doc, ['Hãng', 'RTSP URL Pattern'], [
    ['Hikvision', 'rtsp://user:pass@ip:port/cam/realmonitor?channel=1&subtype=0'],
    ['Dahua', 'rtsp://user:pass@ip:port/cam/realmonitor?channel=1&subtype=0'],
    ['Seetong', 'rtsp://user:pass@ip:port/mpeg4'],
    ['FPT', 'rtsp://user:pass@ip:port/live/0'],
])

doc.add_heading('2.3. Thiết kế cơ sở dữ liệu Camera', level=2)
add_table(doc, ['Field', 'Type', 'Mô tả'], [
    ['id', 'UUID (PK)', 'Khóa chính'],
    ['venue_id', 'UUID (FK)', 'Liên kết venue'],
    ['resource_id', 'UUID (FK)', 'Liên kết sân/bàn cụ thể'],
    ['name', 'VARCHAR(255)', 'Tên camera (VD: Camera Sân 1)'],
    ['ip_address', 'VARCHAR(45)', 'Địa chỉ IP camera'],
    ['port', 'INTEGER', 'Port RTSP (mặc định 554)'],
    ['username', 'VARCHAR(128)', 'Tên đăng nhập camera'],
    ['password', 'VARCHAR(255)', 'Mật khẩu camera'],
    ['camera_brand', 'ENUM', 'Hãng camera (hik/dahua/seetong/fpt/custom)'],
    ['rtsp_url_override', 'VARCHAR(1024)', 'URL RTSP tùy chỉnh (nếu có)'],
    ['is_active', 'BOOLEAN', 'Trạng thái hoạt động'],
])

doc.add_heading('2.4. API Camera', level=2)
add_table(doc, ['Method', 'Endpoint', 'Mô tả', 'Phân quyền'], [
    ['GET', '/api/admin/cameras', 'Admin xem camera theo venue', 'ADMIN'],
    ['POST', '/api/admin/cameras', 'Admin thêm camera', 'ADMIN'],
    ['PATCH', '/api/admin/cameras/{id}', 'Admin sửa camera', 'ADMIN'],
    ['DELETE', '/api/admin/cameras/{id}', 'Admin xoá camera', 'ADMIN'],
    ['GET', '/api/staff/cameras', 'Staff xem camera theo phân công', 'STAFF, ADMIN'],
])

doc.add_heading('2.5. Venue Scope - Phân quyền Camera', level=2)
doc.add_paragraph('Camera được phân quyền theo venue:')
add_table(doc, ['Vai trò', 'Assignment', 'Thấy camera'], [
    ['Admin', 'default_venue_id', 'Chỉ camera trong venue của mình'],
    ['Staff', 'scope=venue (toàn sân)', 'Tất cả camera trong venue'],
    ['Staff', 'scope=resource (sân cụ thể)', 'Chỉ camera gán cho sân đó'],
])

doc.add_heading('2.6. Luồng hoạt động Camera', level=2)
doc.add_paragraph('1. Admin cấu hình camera (IP, port, username, password, hãng)')
doc.add_paragraph('2. Admin gán camera cho sân/bàn cụ thể qua dropdown')
doc.add_paragraph('3. Staff mở "Camera sân" → thấy grid thumbnail các camera được phân công')
doc.add_paragraph('4. Staff chọn camera → media_kit kết nối RTSP → hiển thị video live')

doc.add_heading('2.7. Giao diện Camera', level=2)
doc.add_paragraph('Admin - Quản lý Camera:')
doc.add_paragraph('- CRUD camera với form: tên, IP, port, username, password, hãng, gán sân')
doc.add_paragraph('- Danh sách camera với thông tin: tên, IP, hãng, trạng thái')
doc.add_paragraph()
doc.add_paragraph('Staff - Xem Camera:')
doc.add_paragraph('- Grid 2 cột với thumbnail cards')
doc.add_paragraph('- Mỗi card: icon camera, tên, tên sân, badge LIVE/OFF')
doc.add_paragraph('- Tap để mở full-screen RTSP live stream với indicator LIVE')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PHẦN 3: RESOURCE MANAGEMENT
# ══════════════════════════════════════════════════════════
doc.add_heading('BỔ SUNG 3: QUẢN LÝ SÂN', level=1)
doc.add_paragraph('(Bổ sung vào Chương 3 & 4 - Thiết kế và Triển khai)')

doc.add_heading('3.1. Tổng quan quản lý sân', level=2)
doc.add_paragraph(
    'Tính năng quản lý sân cho phép admin quản lý trạng thái sân/bàn '
    '(hoạt động, bảo trì, tắt) và cấu hình giá theo giờ.'
)

doc.add_heading('3.2. Trạng thái sân', level=2)
add_table(doc, ['Status', 'Ý nghĩa', 'Khách đặt được?', 'Mô tả'], [
    ['active', 'Hoạt động', '✅', 'Sân sẵn sàng cho khách đặt'],
    ['maintenance', 'Bảo trì', '❌', 'Sân đang bảo trì, không thể đặt'],
    ['inactive', 'Tắt', '❌', 'Sân không sử dụng'],
])

doc.add_heading('3.3. API Quản lý sân', level=2)
add_table(doc, ['Method', 'Endpoint', 'Mô tả'], [
    ['GET', '/api/venues/resources', 'Danh sách sân (filter theo status)'],
    ['POST', '/api/admin/resources', 'Admin thêm sân mới'],
    ['PATCH', '/api/admin/resources/{id}', 'Admin cập nhật sân (status, giá, tên)'],
])

doc.add_heading('3.4. Giao diện Quản lý sân', level=2)
doc.add_paragraph('Admin - Quản lý sân:')
doc.add_paragraph('- Filter chips: Tất cả / Hoạt động / Bảo trì / Tắt')
doc.add_paragraph('- Card hiển thị: icon loại sân, tên sân, khu vực, mã sân, giá/giờ')
doc.add_paragraph('- Nút thao tác: Kích hoạt / Bảo trì / Tắt')
doc.add_paragraph('- Badge trạng thái với màu sắc: Xanh (hoạt động), Vàng (bảo trì), Xám (tắt)')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PHẦN 4: BẢNG TỔNG HỢP CHỨC NĂNG
# ══════════════════════════════════════════════════════════
doc.add_heading('BỔ SUNG 4: TỔNG HỢP CHỨC NĂNG HỆ THỐNG', level=1)

doc.add_heading('4.1. Chức năng Customer', level=2)
add_table(doc, ['STT', 'Chức năng', 'Mô tả'], [
    ['1', 'Chat AI', 'Hỏi đáp về luật, kỹ thuật thể thao qua Knowledge Graph'],
    ['2', 'Đặt sân', 'Chọn sân, thời gian, xác nhận đặt với check availability'],
    ['3', 'Đặt món', 'Xem menu, đặt thức ăn/thức uống liên kết booking'],
    ['4', 'Gọi nhân viên', 'Tạo yêu cầu hỗ trợ, thông báo real-time'],
    ['5', 'Xem lịch', 'Lịch sử đặt sân và trạng thái đơn hàng'],
    ['6', 'Thanh toán', 'Stripe (quốc tế) + VNPay (nội địa)'],
    ['7', 'Quét QR nhận sân', 'Quét QR code để xác nhận nhận sân'],
])

doc.add_heading('4.2. Chức năng Staff', level=2)
add_table(doc, ['STT', 'Chức năng', 'Mô tả'], [
    ['1', 'Quản lý đặt sân', 'Xem, xác nhận, check-in booking'],
    ['2', 'Quản lý yêu cầu', 'Tiếp nhận, xử lý yêu cầu khách'],
    ['3', 'Chat với khách', 'Trò chuyện real-time với khách'],
    ['4', 'Xem camera sân', 'Giám sát camera IP các sân được phân công'],
    ['5', 'Quản lý menu', 'Xem, cập nhật trạng thái món'],
    ['6', 'Hoá đơn', 'Xem hoá đơn khách'],
])

doc.add_heading('4.3. Chức năng Admin', level=2)
add_table(doc, ['STT', 'Chức năng', 'Mô tả'], [
    ['1', 'Dashboard', 'Tổng quan doanh thu, booking, orders'],
    ['2', 'Quản lý đặt sân', 'CRUD booking, đổi trạng thái'],
    ['3', 'Quản lý sân', 'Thêm/sửa/xoá sân, đổi trạng thái (active/maintenance/inactive)'],
    ['4', 'Cấu hình giá', 'Thiết lập giá theo giờ cho từng sân'],
    ['5', 'Quản lý camera', 'CRUD camera IP, gán camera cho sân'],
    ['6', 'Quản lý nhân viên', 'CRUD staff, phân quyền theo venue/area/resource'],
    ['7', 'Quản lý menu', 'CRUD thức ăn/thức uống'],
    ['8', 'Biểu đồ', 'Thống kê doanh thu, booking'],
    ['9', 'Quản lý hoá đơn', 'Xem, quản lý hoá đơn'],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PHẦN 5: TÀI LIỆU THAM KHẢO BỔ SUNG
# ══════════════════════════════════════════════════════════
doc.add_heading('TÀI LIỆU THAM KHẢO BỔ SUNG', level=1)
refs = [
    'Flutter Documentation. https://docs.flutter.dev/',
    'Riverpod Documentation. https://riverpod.dev/',
    'GoRouter Package. https://pub.dev/packages/go_router',
    'media_kit Package. https://pub.dev/packages/media_kit',
    'FastAPI Documentation. https://fastapi.tiangolo.com/',
    'SQLAlchemy Documentation. https://docs.sqlalchemy.org/',
    'PostgreSQL Documentation. https://www.postgresql.org/docs/',
    'Neo4j Documentation. https://neo4j.com/docs/',
    'Redis Documentation. https://redis.io/documentation',
    'RTSP Protocol (RFC 2326). https://tools.ietf.org/html/rfc2326',
    'Dio Package. https://pub.dev/packages/dio',
    'flutter_stripe Package. https://pub.dev/packages/flutter_stripe',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'[{i}] {ref}')

# Save merged file
doc.save('Documentations/BaoCaoDoAn_SportsVenueChatbot.docx')
print('Done: Documentations/BaoCaoDoAn_SportsVenueChatbot.docx')
