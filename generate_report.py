from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    return table

# ── TRANG BÌA ──
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TRƯỜNG ĐẠI HỌC ...')
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('KHOA CÔNG NGHỆ THÔNG TIN')
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ĐỒ ÁN TỐT NGHIỆP')
run.font.size = Pt(20)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ĐỀ TÀI:')
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('XÂY DỰNG HỆ THỐNG QUẢN LÝ SÂN THỂ THAO\nTÍCH HỢP AI CHATBOT VÀ GIÁM SÁT CAMERA')
run.font.size = Pt(16)
run.bold = True

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GVHD: ................................')
run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SVTH: ................................')
run.font.size = Pt(13)

doc.add_page_break()

# ── MỤC LỤC ──
doc.add_heading('MỤC LỤC', level=1)
toc_items = [
    'CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI',
    '  1.1. Đặt vấn đề',
    '  1.2. Mục tiêu đề tài',
    '  1.3. Phạm vi nghiên cứu',
    'CHƯƠNG 2: CƠ SỞ LÝ THUYẾT',
    '  2.1. Tổng quan Framework',
    '  2.2. Flutter Framework & Packages',
    '  2.3. FastAPI Backend Framework',
    '  2.4. Hệ quản trị cơ sở dữ liệu',
    '  2.5. Kiến trúc Client-Server',
    '  2.6. Clean Architecture & Repository Pattern',
    '  2.7. Knowledge Graph & Graph RAG',
    '  2.8. AI Agent với Tool Calling',
    '  2.9. Real-time Communication (WebSocket)',
    '  2.10. RTSP Camera Streaming',
    '  2.11. Multi-tenancy',
    '  2.9. Real-time Communication (WebSocket)',
    '  2.10. RTSP Camera Streaming',
    '  2.11. Multi-tenancy',
    'CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ',
    '  3.1. Phân tích yêu cầu',
    '  3.2. Thiết kế cơ sở dữ liệu',
    '  3.3. Thiết kế API',
    '  3.4. Thiết kế giao diện',
    'CHƯƠNG 4: TRIỂN KHAI',
    '  4.1. Công nghệ sử dụng',
    '  4.2. Triển khai Backend',
    '  4.3. Triển khai Frontend',
    '  4.4. Triển khai Camera System',
    'CHƯƠNG 5: KẾT QUẢ VÀ ĐÁNH GIÁ',
    '  5.1. Chức năng đã hoàn thành',
    '  5.2. Demo hệ thống',
    '  5.3. Đánh giá và hạn chế',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# CHƯƠNG 1
# ══════════════════════════════════════════════════════════
doc.add_heading('CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI', level=1)

doc.add_heading('1.1. Đặt vấn đề', level=2)
doc.add_paragraph(
    'Ngày nay, việc quản lý sân thể thao (bida, pickleball, cầu lông) đang trở nên phức tạp hơn '
    'khi số lượng khách hàng tăng trưởng. Các vấn đề thường gặp bao gồm:'
)
issues = [
    'Khách hàng phải gọi điện hoặc đến trực tiếp để đặt sân, gây mất thời gian.',
    'Nhân viên quản lý khó theo dõi trạng thái sân real-time (sân trống, sân đang có khách, sân bảo trì).',
    'Không có hệ thống giám sát camera tập trung, nhân viên phải kiểm tra từng sân.',
    'Khách hàng không có kênh hỗ trợ nhanh khi cần gọi nhân viên hoặc hỏi thông tin.',
    'Quản lý phân quyền nhân viên theo từng sân/khu vực phức tạp.',
]
for issue in issues:
    doc.add_paragraph(issue, style='List Bullet')

doc.add_heading('1.2. Mục tiêu đề tài', level=2)
doc.add_paragraph('Đề tài nhắm đến các mục tiêu sau:')
goals = [
    'Xây dựng hệ thống quản lý sân thể thao đa năng (bida, pickleball, cầu lông).',
    'Tích hợp AI Chatbot sử dụng Knowledge Graph để hỗ trợ khách hàng tự động.',
    'Hệ thống đặt sân trực tuyến với kiểm tra availability real-time.',
    'Tích hợp camera IP (RTSP) để nhân viên giám sát sân từ ứng dụng.',
    'Hệ thống phân quyền nhân viên theo venue/area/resource.',
    'Thanh toán trực tuyến qua Stripe (quốc tế) và VNPay (nội địa).',
    'Thông báo real-time qua WebSocket cho tất cả vai trò.',
]
for goal in goals:
    doc.add_paragraph(goal, style='List Bullet')

doc.add_heading('1.3. Phạm vi nghiên cứu', level=2)
doc.add_paragraph(
    'Đề tài tập trung vào 3 loại sân thể thao: Bida, Pickleball, và Cầu lông. '
    'Hệ thống phục vụ 3 vai trò chính: Customer (khách hàng), Staff (nhân viên), và Admin (quản lý).'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# CHƯƠNG 2
# ══════════════════════════════════════════════════════════
doc.add_heading('CHƯƠNG 2: CƠ SỞ LÝ THUYẾT', level=1)

# ── 2.1 Tổng quan Framework ──
doc.add_heading('2.1. Tổng quan Framework', level=2)
doc.add_paragraph(
    'Framework là bộ khung phần mềm cung cấp các thành phần, công cụ và quy tắc tái sử dụng '
    'để phát triển ứng dụng nhanh hơn, nhất quán hơn và dễ bảo trì hơn. Việc lựa chọn framework '
    'phù hợp giúp giảm thời gian phát triển, tăng tính ổn định và dễ dàng mở rộng hệ thống.'
)
doc.add_paragraph(
    'Hệ thống sử dụng 2 framework chính: Flutter (frontend) và FastAPI (backend), '
    'kết hợp với các thư viện (packages) chuyên biệt cho từng nghiệp vụ.'
)

# ── 2.2 Flutter Framework ──
doc.add_heading('2.2. Flutter Framework & Packages', level=2)

doc.add_heading('2.2.1. Flutter là gì?', level=3)
doc.add_paragraph(
    'Flutter là framework phát triển ứng dụng cross-platform do Google phát triển, '
    'sử dụng ngôn ngữ lập trình Dart. Flutter cho phép xây dựng ứng dụng cho iOS, Android, '
    'Web và Desktop từ một codebase duy nhất.'
)
doc.add_paragraph('Đặc điểm nổi bật của Flutter:')
flutter_features = [
    'Hot Reload: Thay đổi code hiển thị ngay lập tức, tăng tốc development.',
    'Widget-based: Giao diện được xây dựng từ các widget có thể tái sử dụng.',
    'Performance: Render engine riêng (Skia/Impeller), không dùng native UI components.',
    'Cross-platform: Một codebase chạy được trên nhiều nền tảng.',
    'Rich ecosystem: Hàng nghìn packages trên pub.dev.',
]
for f in flutter_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.2.2. Các Packages tiêu biểu', level=3)

doc.add_paragraph('a) State Management - Riverpod:')
doc.add_paragraph(
    'Riverpod là state management solution do tác giả của Provider phát triển, '
    'giải quyết các hạn chế của Provider gốc. Hệ thống sử dụng Riverpod 2.4.9.'
)
doc.add_paragraph('So sánh Riverpod vs Provider vs Bloc:')
add_table(doc, ['Tiêu chí', 'Provider', 'Riverpod', 'Bloc'], [
    ['Compile-safe', '❌ Runtime error', '✅ Compile-time check', '✅ Compile-time check'],
    ['Testability', 'Khó (cần BuildContext)', 'Dễ (independent)', 'Dễ (stream-based)'],
    ['Boilerplate', 'Ít', 'Ít', 'Nhiều (events/states)'],
    ['Learning curve', 'Dễ', 'Trung bình', 'Khó'],
    ['Dependency injection', 'Cần context', 'Global access', 'Manual setup'],
    ['Auto dispose', 'Không', 'Có (autoDispose)', 'Manual'],
    ['Family provider', 'Không', 'Có (family)', 'Không'],
])
doc.add_paragraph(
    'Lựa chọn Riverpod vì: type-safe, không cần BuildContext để đọc state, '
    'hỗ trợ autoDispose và family providers, dễ test.'
)

doc.add_paragraph('b) Navigation - GoRouter:')
doc.add_paragraph(
    'GoRouter là package declarative routing do Flutter team phát triển. '
    'Hỗ trợ deep linking, redirect theo role, nested routes.'
)

doc.add_paragraph('c) Network - Dio:')
doc.add_paragraph(
    'Dio là HTTP client mạnh mẽ cho Dart, hỗ trợ interceptors (auth token, logging), '
    'request cancellation, file upload/download, timeout configuration.'
)

doc.add_paragraph('d) Video Streaming - media_kit:')
doc.add_paragraph(
    'media_kit là package phát video dựa trên libmpv (bản nhẹ của mpv player), '
    'hỗ trợ RTSP, HLS, HTTP streams với hardware-accelerated decoding. '
    'Sử dụng để phát camera RTSP trực tiếp trong ứng dụng.'
)

doc.add_paragraph('e) Các packages khác:')
add_table(doc, ['Package', 'Vai trò'], [
    ['flutter_riverpod', 'State management declarative'],
    ['go_router', 'Declarative navigation với role-based redirect'],
    ['dio', 'HTTP client với interceptors'],
    ['web_socket_channel', 'WebSocket real-time communication'],
    ['flutter_chat_ui', 'Chat UI components'],
    ['flutter_stripe', 'Stripe native payment integration'],
    ['media_kit + media_kit_video', 'RTSP/HLS video playback (libmpv)'],
    ['qr_flutter + mobile_scanner', 'QR code generation & scanning'],
    ['flutter_secure_storage', 'Encrypted storage cho tokens'],
    ['local_auth', 'Biometric/PIN authentication'],
    ['flutter_webrtc', 'Voice call giữa staff và customer'],
])

# ── 2.3 FastAPI Backend ──
doc.add_heading('2.3. FastAPI Backend Framework', level=2)

doc.add_heading('2.3.1. FastAPI là gì?', level=3)
doc.add_paragraph(
    'FastAPI là framework xây dựng API hiện đại bằng Python, '
    'tốc độ ngang với Node.js và Go nhờ sử dụng ASGI (Asynchronous Server Gateway Interface).'
)
doc.add_paragraph('Đặc điểm nổi bật:')
fastapi_features = [
    'Tốc độ cao: Sử dụng Starlette và Pydantic, async/await native.',
    'Type-safe: Tự động validate request/response dựa trên type hints.',
    'Auto documentation: Tự sinh Swagger UI và ReDoc từ code.',
    'Dependency Injection: Hệ thống dependency injection mạnh mẽ.',
    'Async support: Hỗ trợ asyncio cho I/O-bound operations.',
]
for f in fastapi_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.3.2. SQLAlchemy ORM', level=3)
doc.add_paragraph(
    'SQLAlchemy là ORM (Object-Relational Mapping) phổ biến nhất cho Python, '
    'cho phép tương tác với database bằng Python objects thay vì raw SQL.'
)
doc.add_paragraph('Hệ thống sử dụng SQLAlchemy 2.0 với async mode (asyncpg driver):')
sa_features = [
    'Declarative models: Định nghĩa bảng bằng Python classes.',
    'Async session: Sử dụng asyncpg cho PostgreSQL async queries.',
    'Relationship mapping: Tự động join và load related objects.',
    'Migration support: Hỗ trợ schema migration qua Alembic.',
]
for f in sa_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('2.3.3. Pydantic Validation', level=3)
doc.add_paragraph(
    'Pydantic là thư viện validation dữ liệu dựa trên Python type annotations, '
    'tự động serialize/deserialize JSON, và sinh JSON Schema.'
)

# ── 2.4 Hệ quản trị CSDL ──
doc.add_heading('2.4. Hệ quản trị cơ sở dữ liệu', level=2)
doc.add_paragraph(
    'Hệ thống sử dụng 3 loại cơ sở dữ liệu khác nhau, mỗi loại phục vụ mục đích riêng biệt:'
)

doc.add_heading('2.4.1. PostgreSQL (Relational Database)', level=3)
doc.add_paragraph(
    'PostgreSQL là hệ quản trị cơ sở dữ liệu quan hệ (RDBMS) mã nguồn mở, '
    'hỗ trợ ACID transactions, JSON data type, và full-text search.'
)
doc.add_paragraph('Vai trò trong hệ thống:')
pg_roles = [
    'Lưu trữ dữ liệu chính: users, bookings, orders, venues, resources, cameras.',
    'ACID compliance: Đảm bảo tính nhất quán dữ liệu cho giao dịch.',
    'Foreign keys: Duy trì tính toàn vẹn tham chiếu giữa các bảng.',
    'JSONB: Lưu trữ dữ liệu linh hoạt (metadata, payload).',
    'UUID primary keys: Phân tán dữ liệu an toàn.',
]
for r in pg_roles:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('2.4.2. Neo4j (Graph Database)', level=3)
doc.add_paragraph(
    'Neo4j là hệ quản trị cơ sở dữ liệu đồ thị (Graph DBMS), '
    'lưu trữ dữ liệu dưới dạng nodes (đỉnh) và relationships (cạnh).'
)
doc.add_paragraph('So sánh Relational vs Graph Database:')
add_table(doc, ['Tiêu chí', 'PostgreSQL (Relational)', 'Neo4j (Graph)'], [
    ['Cấu trúc', 'Bảng (rows/columns)', 'Đồ thị (nodes/relationships)'],
    ['Truy vấn quan hệ', 'JOIN (chậm khi phức tạp)', 'Pattern matching (nhanh)'],
    ['Schema', 'Cố định (schema-on-write)', 'Linh hoạt (schema-free)'],
    ['Use case', 'OLTP, báo cáo', 'Knowledge graph, social network, recommendation'],
    ['Ngôn ngữ truy vấn', 'SQL', 'Cypher'],
])
doc.add_paragraph('Vai trò trong hệ thống:')
neo4j_roles = [
    'Lưu trữ tri thức thể thao: luật chơi, kỹ thuật, thiết bị.',
    'Graph RAG: Truy vấn tri thức kết hợp LLM cho AI chatbot.',
    'Full-text search: Tìm kiếm nhanh theo nội dung.',
]
for r in neo4j_roles:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('2.4.3. Redis (In-memory Database)', level=3)
doc.add_paragraph(
    'Redis là hệ quản trị cơ sở dữ liệu in-memory, '
    'hỗ trợ nhiều cấu trúc dữ liệu (string, list, set, hash, stream).'
)
doc.add_paragraph('Vai trò trong hệ thống:')
redis_roles = [
    'Session storage: Lưu trữ chat session với TTL (time-to-live).',
    'Pub/Sub: Publish/subscribe cho real-time notifications.',
    'Ephemeral chat rooms: Lưu tin nhắn chat staff-customer (24h TTL).',
    'Presence tracking: Theo dõi trạng thái online/offline.',
    'Caching: Cache dữ liệu thường truy cập.',
]
for r in redis_roles:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('2.4.4. So sánh 3 loại Database', level=3)
add_table(doc, ['Tiêu chí', 'PostgreSQL', 'Neo4j', 'Redis'], [
    ['Loại', 'Relational', 'Graph', 'In-memory'],
    ['Storage', 'Disk', 'Disk', 'RAM (persistent optional)'],
    ['Speed', 'Nhanh (indexed)', 'Nhanh (graph traversal)', 'Rất nhanh (in-memory)'],
    ['ACID', '✅ Có', '✅ Có', '✅ Có (Lua scripting)'],
    ['Use case chính', 'OLTP, CRUD', 'Knowledge graph', 'Cache, pub/sub, session'],
    ['Scalability', 'Vertical + Read replica', 'Vertical + Cluster', 'Horizontal (cluster)'],
    ['Ngôn ngữ', 'SQL', 'Cypher', 'Redis commands'],
])

# ── 2.5 Client-Server ──
doc.add_heading('2.5. Kiến trúc Client-Server', level=2)
doc.add_paragraph(
    'Kiến trúc Client-Server là mô hình trong đó client gửi yêu cầu đến server, '
    'server xử lý và trả kết quả về. Hệ thống áp dụng mô hình 3-tier:'
)
tiers = [
    'Presentation Layer (Client): Flutter mobile app với giao diện cho Customer, Staff, Admin.',
    'Application Layer (Server): FastAPI backend xử lý business logic, AI agent, camera streaming.',
    'Data Layer: PostgreSQL (dữ liệu chính), Neo4j (knowledge graph), Redis (cache, pub/sub).',
]
for t in tiers:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('2.6. Clean Architecture & Repository Pattern', level=2)
doc.add_paragraph(
    'Clean Architecture là nguyên tắc thiết kế phần mềm do Robert C. Martin đề xuất, '
    'tách biệt các tầng trách nhiệm:'
)
layers = [
    'API Layer: Xử lý HTTP request/response, validation.',
    'Service Layer: Chứa business logic, điều phối các repository.',
    'Repository Layer: Truy cập dữ liệu,抽象 hoá database operations.',
    'Model Layer: Định nghĩa cấu trúc dữ liệu (SQLAlchemy ORM).',
    'Schema Layer: Validation dữ liệu đầu vào/ra (Pydantic).',
]
for l in layers:
    doc.add_paragraph(l, style='List Bullet')

doc.add_heading('2.7. Knowledge Graph & Graph RAG', level=2)
doc.add_paragraph(
    'Knowledge Graph (Đồ thị tri thức) là cấu trúc dữ liệu biểu diễn tri thức dưới dạng đồ thị, '
    'trong đó Node đại diện cho thực thể và Edge đại diện cho mối quan hệ.'
)
doc.add_paragraph(
    'Graph RAG (Retrieval-Augmented Generation) kết hợp truy vấn đồ thị tri thức với LLM '
    'để tạo câu trả lời chính xác, giảm hallucination bằng cách grounding vào dữ liệu thực.'
)
doc.add_paragraph('Schema đồ thị tri thức:')
add_table(doc, ['Loại', 'Ví dụ'], [
    ['Node types', 'Rule, Technique, Equipment, Sport, Concept, GameType'],
    ['Relationships', 'DUNG_DE, LIEN_QUAN, LA_LOAI, THUOC, SU_DUNG, QUY_DINH'],
    ['Stats', '418 entities, 441 relationships, full-text search index'],
])

doc.add_heading('2.8. AI Agent với Tool Calling', level=2)
doc.add_paragraph(
    'AI Agent là hệ thống AI có khả năng nhận diện ý định từ câu hỏi, chọn tool phù hợp, '
    'và thực thi tool để hoàn thành tác vụ. Hệ thống sử dụng LangChain framework.'
)
add_table(doc, ['Tool', 'Trigger', 'Chức năng'], [
    ['query_knowledge', 'Hỏi luật, kỹ thuật', 'Truy vấn Neo4j knowledge graph'],
    ['book_court', 'Đặt sân', 'Kiểm tra availability, tạo booking'],
    ['order_food', 'Gọi đồ', 'Tạo order, liên kết booking'],
    ['call_staff', 'Gọi nhân viên', 'Tạo staff request, thông báo WebSocket'],
    ['check_schedule', 'Xem lịch', 'Query bookings theo ngày'],
    ['order_menu_items', 'Đặt món', 'Đặt từ menu với booking association'],
])

doc.add_heading('2.9. Real-time Communication (WebSocket)', level=2)
doc.add_paragraph(
    'WebSocket là giao thức truyền thông hai chiều (full-duplex) trên một kết nối TCP, '
    'cho phép server gửi dữ liệu đến client mà không cần client yêu cầu.'
)
doc.add_paragraph('Áp dụng trong hệ thống:')
ws_uses = [
    'Notifications: Thông báo real-time cho staff/admin khi có sự kiện (booking, order, payment).',
    'Staff Chat: Chat trực tiếp giữa staff và customer với room-scoped messaging.',
    'Court Status: Cập nhật trạng thái sân real-time (checked_in, completed).',
    'Call Signaling: WebRTC signaling cho voice call giữa staff và customer.',
]
for use in ws_uses:
    doc.add_paragraph(use, style='List Bullet')

doc.add_heading('2.10. RTSP Camera Streaming', level=2)
doc.add_paragraph(
    'RTSP (Real Time Streaming Protocol) là giao thức truyền phát video trực tuyến, '
    'được sử dụng phổ biến trong hệ thống camera IP.'
)
doc.add_paragraph('Kiến trúc streaming:')
doc.add_paragraph('[IP Camera] --RTSP TCP--> [media_kit/libmpv] --> [Flutter Video Widget]')
doc.add_paragraph('Hỗ trợ các hãng camera:')
add_table(doc, ['Hãng', 'RTSP URL Pattern'], [
    ['Hikvision', 'rtsp://user:pass@ip:port/cam/realmonitor?channel=1&subtype=0'],
    ['Dahua', 'rtsp://user:pass@ip:port/cam/realmonitor?channel=1&subtype=0'],
    ['Seetong', 'rtsp://user:pass@ip:port/mpeg4'],
    ['FPT', 'rtsp://user:pass@ip:port/live/0'],
])
doc.add_paragraph(
    'media_kit sử dụng libmpv (bản nhẹ của mpv player) để giải mã video hardware-accelerated '
    'trên Android, hỗ trợ H.264/H.265 codec với độ trễ thấp.'
)

doc.add_heading('2.11. Multi-tenancy', level=2)
doc.add_paragraph(
    'Multi-tenancy là kiến trúc một ứng dụng phục vụ nhiều khách hàng (tenant) khác nhau, '
    'mỗi tenant có dữ liệu và cấu hình riêng biệt.'
)
doc.add_paragraph('Triển khai trong hệ thống:')
tenancy_items = [
    'Business: Tenant chính, mỗi business có nhiều venue.',
    'Venue: Địa điểm thể thao (CLB Bida, Sân Pickleball, Nhà thi đấu).',
    'ServiceResource: Sân/bàn cụ thể trong venue.',
    'StaffAssignment: Liên kết staff với venue/area/resource với scope-based access control.',
]
for item in tenancy_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# CHƯƠNG 3
# ══════════════════════════════════════════════════════════
doc.add_heading('CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ', level=1)

doc.add_heading('3.1. Phân tích yêu cầu', level=2)
doc.add_heading('3.1.1. Yêu cầu chức năng', level=3)

doc.add_paragraph('a) Customer:')
cust_reqs = [
    'Chat AI hỏi đáp về luật, kỹ thuật thể thao.',
    'Đặt sân trực tuyến với kiểm tra availability.',
    'Đặt thức ăn/thức uống, liên kết với booking.',
    'Gọi nhân viên hỗ trợ.',
    'Thanh toán trực tuyến (Stripe/VNPay).',
    'Quét QR code nhận sân.',
]
for r in cust_reqs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph('b) Staff:')
staff_reqs = [
    'Quản lý đặt sân (xem, xác nhận, check-in).',
    'Tiếp nhận và xử lý yêu cầu khách.',
    'Chat real-time với khách.',
    'Giám sát camera IP các sân được phân công.',
    'Quản lý hoá đơn.',
]
for r in staff_reqs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph('c) Admin:')
admin_reqs = [
    'Dashboard tổng quan doanh thu, booking, orders.',
    'Quản lý sân (CRUD, đổi trạng thái active/maintenance/inactive).',
    'Cấu hình giá sân theo giờ.',
    'Quản lý camera IP (CRUD, gán camera cho sân).',
    'Quản lý nhân viên và phân quyền.',
    'Quản lý menu thức ăn/thức uống.',
    'Biểu đồ thống kê doanh thu.',
]
for r in admin_reqs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('3.2. Thiết kế cơ sở dữ liệu', level=2)
doc.add_paragraph('Schema chính (PostgreSQL):')

add_table(doc, ['Bảng', 'Mô tả', 'Fields chính'], [
    ['users', 'Tài khoản người dùng', 'id, phone, name, role, default_venue_id'],
    ['venues', 'Địa điểm thể thao', 'id, business_id, name, address'],
    ['service_resources', 'Sân/bàn', 'id, venue_id, code, name, status, hourly_rate'],
    ['cameras', 'Camera IP', 'id, venue_id, resource_id, ip_address, camera_brand'],
    ['bookings', 'Đặt sân', 'id, user_id, resource_id, start_time, status'],
    ['orders', 'Đặt món', 'id, booking_id, venue_id, total_amount, payment_status'],
    ['staff_assignments', 'Phân quyền', 'id, staff_id, venue_id, scope'],
    ['staff_requests', 'Yêu cầu hỗ trợ', 'id, user_id, venue_id, request_type, status'],
    ['notifications', 'Thông báo', 'id, event_type, title, target_roles'],
])

doc.add_heading('3.3. Thiết kế API', level=2)
doc.add_paragraph('Hệ thống cung cấp hơn 70 REST endpoints và 2 WebSocket endpoints:')
add_table(doc, ['Nhóm', 'Số lượng', 'Mô tả'], [
    ['Auth', '4', 'Đăng nhập, xác thực, đổi mật khẩu'],
    ['Chat', '1', 'AI chat với tool calling'],
    ['Booking', '10', 'CRUD, availability, check-in, bills'],
    ['Order', '4', 'CRUD order với booking linking'],
    ['Menu', '3', 'List, top-selling, suggest'],
    ['Staff Request', '7', 'Create, accept, complete, cancel'],
    ['Staff Chat', '4', 'Rooms, history, close, WebSocket'],
    ['Realtime', '4', 'WebSocket notifications, list, mark read'],
    ['Payment', '9', 'Stripe + VNPay endpoints'],
    ['Admin', '15', 'Dashboard, bookings, cameras, staff'],
    ['Venue', '8+', 'CRUD venues, resources, cameras'],
])

doc.add_heading('3.4. Thiết kế giao diện', level=2)
doc.add_paragraph('Giao diện Flutter với 3 vai trò:')
roles_ui = [
    'Customer: Home (Chat AI), Đặt sân, Menu, Hoá đơn, Hồ sơ.',
    'Staff: Đặt sân, Yêu cầu, Thực đơn, Hoá đơn, Tin nhắn + Camera (popup menu).',
    'Admin: Dashboard, Đặt sân, Thực đơn, Nhân viên, Hoá đơn + Quản lý sân, Camera (popup menu).',
]
for r in roles_ui:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# CHƯƠNG 4
# ══════════════════════════════════════════════════════════
doc.add_heading('CHƯƠNG 4: TRIỂN KHAI', level=1)

doc.add_heading('4.1. Công nghệ sử dụng', level=2)
add_table(doc, ['Layer', 'Công nghệ', 'Vai trò'], [
    ['Mobile App', 'Flutter 3.1+ / Dart', 'Cross-platform UI cho iOS/Android'],
    ['Backend API', 'Python / FastAPI 0.115+', 'REST API server bất đồng bộ'],
    ['AI Agent', 'LangChain 0.3.1 + Ollama', 'Chat AI với tool calling'],
    ['LLM', 'qwen2.5-coder:7b', 'Xử lý ngôn ngữ tự nhiên (local)'],
    ['Knowledge Graph', 'Neo4j AuraDB 5.x', 'Lưu trữ tri thức thể thao'],
    ['Database', 'PostgreSQL 14+', 'Cơ sở dữ liệu chính (asyncpg)'],
    ['Cache/PubSub', 'Redis 7+', 'Session, WebSocket pub/sub, chat'],
    ['Video Streaming', 'media_kit 1.1.10+ (libmpv)', 'RTSP camera playback'],
    ['Payment', 'Stripe + VNPay', 'Thanh toán quốc tế + nội địa'],
    ['Real-time', 'WebSocket (web_socket_channel)', 'Thông báo real-time'],
])

doc.add_heading('4.2. Triển khai Backend', level=2)
doc.add_paragraph('Backend được xây dựng với FastAPI, áp dụng Clean Architecture:')
be_components = [
    'API Layer (app/api/): 15 FastAPI routers xử lý HTTP requests.',
    'Service Layer (app/services/): Business logic cho booking, chat, notification, payment.',
    'Repository Layer (app/repositories/): Data access pattern, truy vấn PostgreSQL.',
    'Model Layer (app/models/): SQLAlchemy ORM với UUID primary keys.',
    'Schema Layer (app/schemas/): Pydantic validation với 59+ schema classes.',
    'Agent Layer (app/agent/): LangChain AI agent với 6 tools.',
]
for c in be_components:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('4.3. Triển khai Frontend', level=2)
doc.add_paragraph('Flutter app với Riverpod state management:')
fe_components = [
    '13 feature modules: auth, chat, booking, menu, payment, staff, admin, camera, ...',
    'GoRouter cho declarative routing với role-based redirect.',
    'Dio client với auth interceptor cho HTTP requests.',
    'WebSocket connections cho real-time notifications và staff chat.',
    'media_kit (libmpv) cho RTSP camera playback.',
    'flutter_stripe cho native payment integration.',
]
for c in fe_components:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('4.4. Triển khai Camera System', level=2)
doc.add_paragraph('Hệ thống camera RTSP tích hợp:')
cam_details = [
    'Backend: Camera model (IP, port, username, password, brand), CRUD API, venue-scoped access.',
    'Admin UI: Quản lý camera (thêm/sửa/xoá), gán camera cho sân/bàn qua dropdown.',
    'Staff UI: Grid thumbnail camera, tap để mở full-screen RTSP live stream.',
    'RTSP Playback: media_kit sử dụng libmpv với hardware-accelerated decoding.',
    'Access Control: Admin chỉ quản lý camera trong venue, staff chỉ thấy camera sân được phân công.',
]
for d in cam_details:
    doc.add_paragraph(d, style='List Bullet')

doc.add_paragraph('Luồng camera:')
doc.add_paragraph('1. Admin thêm camera (IP, port, username, password, hãng camera)')
doc.add_paragraph('2. Admin gán camera cho sân/bàn cụ thể qua dropdown')
doc.add_paragraph('3. Staff mở "Camera sân" → thấy grid thumbnail các camera được phân công')
doc.add_paragraph('4. Staff chọn camera → media_kit kết nối RTSP → hiển thị video live')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# CHƯƠNG 5
# ══════════════════════════════════════════════════════════
doc.add_heading('CHƯƠNG 5: KẾT QUẢ VÀ ĐÁNH GIÁ', level=1)

doc.add_heading('5.1. Chức năng đã hoàn thành', level=2)
add_table(doc, ['Chức năng', 'Trạng thái', 'Mô tả'], [
    ['AI Chatbot', '✅ Hoàn thành', 'Chat AI với Knowledge Graph, 6 tools'],
    ['Đặt sân', '✅ Hoàn thành', 'Đặt sân online, check availability, QR check-in'],
    ['Đặt món', '✅ Hoàn thành', 'Menu CRUD, order với booking linking'],
    ['Thanh toán', '✅ Hoàn thành', 'Stripe + VNPay, bill aggregation'],
    ['Staff Request', '✅ Hoàn thành', 'Create, accept, complete, cancel'],
    ['Staff Chat', '✅ Hoàn thành', 'Real-time chat, WebSocket, ephemeral rooms'],
    ['Notifications', '✅ Hoàn thành', 'WebSocket push, role-based targeting'],
    ['Admin Dashboard', '✅ Hoàn thành', 'Stats, analytics, recent activity'],
    ['Quản lý sân', '✅ Hoàn thành', 'CRUD, status toggle (active/maintenance/inactive)'],
    ['Cấu hình giá', '✅ Hoàn thành', 'Hourly rate per resource'],
    ['Camera System', '✅ Hoàn thành', 'RTSP streaming, venue-scoped, admin CRUD'],
    ['Phân quyền', '✅ Hoàn thành', 'Staff assignment với venue/area/resource scope'],
])

doc.add_heading('5.2. Demo hệ thống', level=2)
doc.add_paragraph('Tài khoản demo:')
add_table(doc, ['Vai trò', 'Tên', 'SĐT', 'Mật khẩu', 'Venue'], [
    ['Admin Bida', 'Quản lý Bida', '0111111111', '123456', 'CLB Bida Sài Gòn'],
    ['Staff Bida', 'NV Bida', '0111111112', '123456', 'CLB Bida Sài Gòn'],
    ['Admin Pickleball', 'Quản lý Pickleball', '0222222222', '123456', 'Sân Pickleball Thủ Đức'],
    ['Staff Pickleball', 'NV Pickleball', '0222222223', '123456', 'Sân Pickleball Thủ Đức'],
    ['Admin Cầu lông', 'Quản lý Cầu lông', '0333333333', '123456', 'Nhà thi đấu Cầu lông Bình Thạnh'],
    ['Khách hàng', 'Khách hàng', '0900000000', '123456', '(mặc định: CLB Bida)'],
])

doc.add_heading('5.3. Đánh giá và hạn chế', level=2)
doc.add_paragraph('Ưu điểm:')
pros = [
    'Hệ thống đa năng, hỗ trợ nhiều loại sân thể thao.',
    'AI Chatbot với Knowledge Graph cho câu trả lời chính xác.',
    'Camera RTSP tích hợp trực tiếp trong app, không cần phần mềm riêng.',
    'Phân quyền linh hoạt theo venue/area/resource.',
    'Thanh toán đa phương thức (Stripe + VNPay).',
]
for p_item in pros:
    doc.add_paragraph(p_item, style='List Bullet')

doc.add_paragraph('Hạn chế:')
cons = [
    'Camera chỉ hoạt động trong mạng LAN (cùng mạng với camera IP).',
    'Chưa hỗ trợ cloud streaming cho camera từ xa.',
    'LLM local (Ollama) có thể chậm trên máy cấu hình thấp.',
    'Chưa có push notification native (FCM/APNs).',
]
for c_item in cons:
    doc.add_paragraph(c_item, style='List Bullet')

doc.add_paragraph('Hướng phát triển:')
future = [
    'Thêm cloud streaming proxy cho camera (RTSP → HLS/WebSocket).',
    'Tích hợp push notification (Firebase Cloud Messaging).',
    'Mở rộng AI Agent với nhiều tool hơn (thống kê, báo cáo).',
    'Deploy lên cloud (AWS/GCP) cho production.',
]
for f_item in future:
    doc.add_paragraph(f_item, style='List Bullet')

doc.add_page_break()

doc.add_heading('TÀI LIỆU THAM KHẢO', level=1)
refs = [
    'FastAPI Documentation. https://fastapi.tiangolo.com/',
    'Flutter Documentation. https://docs.flutter.dev/',
    'SQLAlchemy Documentation. https://docs.sqlalchemy.org/',
    'Neo4j Documentation. https://neo4j.com/docs/',
    'LangChain Documentation. https://python.langchain.com/',
    'RTSP Protocol (RFC 2326). https://tools.ietf.org/html/rfc2326',
    'media_kit Package. https://pub.dev/packages/media_kit',
    'Riverpod Documentation. https://riverpod.dev/',
    'Stripe API Documentation. https://stripe.com/docs/api',
    'VNPay Developer Documentation. https://sandbox.vnpayment.vn/apis/',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'[{i}] {ref}')

doc.save('Documentations/BaoCaoDoAn_SportsVenueChatbot.docx')
print('Done: Documentations/BaoCaoDoAn_SportsVenueChatbot.docx')
