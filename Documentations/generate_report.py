#!/usr/bin/env python3
"""Generate graduation report (.docx) following university formatting guidelines."""

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import os


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_format(paragraph, space_before=0, space_after=0, line_spacing=1.5):
    """Set paragraph formatting."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing


def add_heading_custom(doc, text, level=1):
    """Add heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = None
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
    set_paragraph_format(heading, space_before=12, space_after=6)
    return heading


def add_body_text(doc, text, bold=False, alignment=None):
    """Add body text with proper formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.bold = bold
    set_paragraph_format(p, line_spacing=1.5)
    if alignment:
        p.alignment = alignment
    return p


def add_bullet_point(doc, text):
    """Add bullet point with proper formatting."""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    set_paragraph_format(p, line_spacing=1.5)
    return p


def create_table_with_header(doc, headers, data, col_widths=None):
    """Create table with header row formatting."""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            set_paragraph_format(paragraph, line_spacing=1.15)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                set_paragraph_format(paragraph, line_spacing=1.15)

    return table


def add_page_break(doc):
    """Add page break."""
    doc.add_page_break()


def add_mermaid_block(doc, title, mermaid_code):
    """Add a Mermaid diagram code block with title."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    set_paragraph_format(p, space_before=6, space_after=4)

    p2 = doc.add_paragraph()
    run_label = p2.add_run("Mermaid code (copy để generate hình):")
    run_label.font.name = "Times New Roman"
    run_label.font.size = Pt(11)
    run_label.italic = True

    # Code block
    code_p = doc.add_paragraph()
    code_run = code_p.add_run(mermaid_code)
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(9)
    pf = code_p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Add shading to code block
    from docx.oxml.ns import qn as _qn
    from docx.oxml import parse_xml as _parse_xml
    shading = _parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    code_p.paragraph_format.element.get_or_add_pPr().append(shading)


def create_report():
    doc = Document()

    # ===== PAGE SETUP =====
    # Set margins: Left 3.5cm, Right 2cm, Top 2.5cm, Bottom 2.5cm
    for section in doc.sections:
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # ===== DEFAULT FONT =====
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.5

    # ====================================================================
    # TRANG BÌA
    # ====================================================================
    for _ in range(3):
        doc.add_paragraph()

    # Logo placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[LOGO TRƯỜNG]")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC ................")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐỒ ÁN TỐT NGHIỆP")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐỀ TÀI:")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("XÂY DỰNG CHATBOT AI CHO QUÁN THỂ THAO")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TÍCH HỢP KNOWLEDGE GRAPH VÀ TOOL CALLING")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    for _ in range(4):
        doc.add_paragraph()

    # Thông tin sinh viên
    info_items = [
        ("Sinh viên thực hiện:", "........................"),
        ("Mã số sinh viên:", "........................"),
        ("Lớp:", "........................"),
        ("Giáo viên hướng dẫn:", "........................"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label} {value}")
        run.font.size = Pt(13)
        run.font.name = "Times New Roman"

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TP. Hồ Chí Minh, năm 2025")
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"
    run.italic = True

    add_page_break(doc)

    # ====================================================================
    # TRANG BÌA TRONG (giống trang bìa)
    # ====================================================================
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG ĐẠI HỌC ................")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐỒ ÁN TỐT NGHIỆP")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐỀ TÀI:")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("XÂY DỰNG CHATBOT AI CHO QUÁN THỂ THAO")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TÍCH HỢP KNOWLEDGE GRAPH VÀ TOOL CALLING")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    for _ in range(4):
        doc.add_paragraph()

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label} {value}")
        run.font.size = Pt(13)
        run.font.name = "Times New Roman"

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TP. Hồ Chí Minh, năm 2025")
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"
    run.italic = True

    add_page_break(doc)

    # ====================================================================
    # LỜI MỞ ĐẦU (không đánh số trang)
    # ====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LỜI MỞ ĐẦU")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    add_body_text(doc,
        "Trong bối cảnh công nghệ trí tuệ nhân tạo (AI) đang phát triển mạnh mẽ, "
        "việc ứng dụng AI vào quản lý và phục vụ khách hàng trong lĩnh vực thể thao "
        "giải trí ngày càng trở nên cần thiết. Đồ án này tập trung vào việc xây dựng "
        "hệ thống chatbot AI cho quán thể thao (bida, pickleball, cầu lông) "
        "tích hợp Knowledge Graph và Tool Calling."
    )

    add_body_text(doc,
        "Lý do chọn đề tài: Các quán thể thao hiện nay vẫn chủ yếu quản lý "
        "theo cách thủ công, thiếu hệ thống tự động hóa để hỗ trợ khách hàng "
        "tra cứu luật chơi, kỹ thuật, đặt sân và gọi đồ uống. Việc xây dựng "
        "chatbot AI giúp giải quyết các vấn đề trên, nâng cao trải nghiệm "
        "khách hàng và tối ưu hóa quy trình phục vụ. Bên cạnh đó, nhu cầu "
        "thanh toán trực tuyến ngày càng cao, đòi hỏi hệ thống cần tích hợp "
        "các cổng thanh toán phổ biến như Stripe và VNPay."
    )

    add_body_text(doc,
        "Mục tiêu nghiên cứu: Xây dựng hệ thống chatbot AI có khả năng "
        "(1) trả lời câu hỏi về luật chơi, kỹ thuật thể thao thông qua Knowledge Graph; "
        "(2) hỗ trợ đặt sân theo thời gian thực; "
        "(3) đặt đồ uống từ menu; "
        "(4) gọi nhân viên hỗ trợ; "
        "(5) tích hợp thanh toán trực tuyến (Stripe, VNPay); "
        "(6) cung cấp ứng dụng di động đa nền tảng."
    )

    add_body_text(doc,
        "Phương pháp nghiên cứu: Đồ án sử dụng phương pháp kết hợp giữa "
        "phát triển phần mềm theo mô hình Agile và ứng dụng các công nghệ AI "
        "hiện đại (LangChain, LLM, Knowledge Graph). Dữ liệu tri thức thể thao "
        "được thu thập từ các nguồn chính thức (WPA, USAPA, BWF) và xây dựng "
        "thành Knowledge Graph trong Neo4j. Hệ thống được triển khai bằng "
        "Docker Compose với kiến trúc microservice."
    )

    add_body_text(doc,
        "Kết cấu của đề tài: Đồ án gồm 4 chương: "
        "Chương 1 - Tổng quan; "
        "Chương 2 - Phân tích; "
        "Chương 3 - Thiết kế cơ sở dữ liệu; "
        "Chương 4 - Hiện thực chương trình."
    )

    add_page_break(doc)

    # ====================================================================
    # NHIỆM VỤ THIẾT KẾ (không đánh số trang)
    # ====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NHIỆM VỤ THIẾT KẾ")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    add_body_text(doc, "Đề tài: Xây dựng Chatbot AI cho quán thể thao tích hợp Knowledge Graph và Tool Calling")

    doc.add_paragraph()
    add_body_text(doc, "Nội dung cần thực hiện:", bold=True)

    tasks = [
        "Nghiên cứu các công nghệ AI: LangChain, LLM (Large Language Model), Knowledge Graph.",
        "Nghiên cứu kiến trúc Tool Calling và Graph RAG (Retrieval Augmented Generation).",
        "Thu thập và xây dựng Knowledge Graph tri thức thể thao (bida, pickleball, cầu lông).",
        "Thiết kế và xây dựng Backend API với FastAPI (Python).",
        "Thiết kế và xây dựng ứng dụng di động Flutter đa nền tảng.",
        "Tích hợp AI Agent với các tool: tra cứu kiến thức, đặt sân, gọi đồ uống, gọi nhân viên.",
        "Thiết kế cơ sở dữ liệu PostgreSQL cho dữ liệu nghiệp vụ.",
        "Tích hợp thanh toán trực tuyến: Stripe (quốc tế) và VNPay (trong nước).",
        "Triển khai hệ thống bằng Docker Compose.",
        "Thử nghiệm và đánh giá hệ thống.",
    ]
    for task in tasks:
        add_bullet_point(doc, task)

    doc.add_paragraph()
    add_body_text(doc, "Yêu cầu về sản phẩm:", bold=True)

    product_reqs = [
        "Hệ thống chatbot AI hoạt động, trả lời được câu hỏi về luật chơi và kỹ thuật thể thao.",
        "Chức năng đặt sân thể thao hoạt động theo thời gian thực.",
        "Chức năng đặt đồ uống từ menu hoạt động.",
        "Chức năng thanh toán trực tuyến hoạt động (Stripe và VNPay).",
        "Ứng dụng di động Flutter chạy được trên Android và iOS.",
        "Hệ thống triển khai được bằng Docker Compose.",
        "Tài liệu mã nguồn và báo cáo đồ án hoàn chỉnh.",
    ]
    for req in product_reqs:
        add_bullet_point(doc, req)

    doc.add_paragraph()
    doc.add_paragraph()

    # Chữ ký GVHD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("TP. Hồ Chí Minh, ngày .... tháng .... năm 2025")
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Remove cell borders
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                border.set(qn("w:sz"), "0")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")
                tcBorders.append(border)
            tcPr.append(tcBorders)

    cell1 = table.rows[0].cells[0]
    p = cell1.paragraphs[0]
    run = p.add_run("GIÁO VIÊN HƯỚNG DẪN")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    cell2 = table.rows[0].cells[1]
    p = cell2.paragraphs[0]
    run = p.add_run("SINH VIÊN THỰC HIỆN")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"

    cell3 = table.rows[1].cells[0]
    p = cell3.paragraphs[0]
    run = p.add_run("(Ký và ghi rõ họ tên)")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.italic = True

    cell4 = table.rows[1].cells[1]
    p = cell4.paragraphs[0]
    run = p.add_run("(Ký và ghi rõ họ tên)")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.italic = True

    add_page_break(doc)

    # ====================================================================
    # MỤC LỤC (không đánh số trang)
    # ====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MỤC LỤC")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    toc_entries = [
        ("LỜI MỞ ĐẦU", "i"),
        ("NHIỆM VỤ THIẾT KẾ", "ii"),
        ("MỤC LỤC", "iii"),
        ("DANH MỤC CÁC BẢNG BIỂU", "iv"),
        ("BẢNG TỪ VIẾT TẮT", "v"),
        ("CHƯƠNG 1. TỔNG QUAN", "1"),
        ("1.1. Đặt vấn đề", "1"),
        ("1.2. Nhiệm vụ của đồ án", "2"),
        ("1.2.1. Mục đích", "2"),
        ("1.2.2. Yêu cầu", "3"),
        ("1.2.3. Môi trường phát triển", "4"),
        ("1.3. Giới thiệu công nghệ", "5"),
        ("1.3.1. Flutter", "5"),
        ("1.3.2. FastAPI (Python)", "6"),
        ("1.3.3. LangChain + Ollama (AI Agent)", "7"),
        ("1.3.4. Neo4j (Knowledge Graph)", "8"),
        ("1.3.5. PostgreSQL", "9"),
        ("1.3.6. Redis", "10"),
        ("1.3.7. Stripe", "11"),
        ("1.3.8. VNPay", "12"),
        ("1.3.9. Docker", "13"),
        ("CHƯƠNG 2. PHÂN TÍCH", "14"),
        ("2.1. Xác định yêu cầu", "14"),
        ("2.1.1. Yêu cầu chức năng", "14"),
        ("2.1.2. Yêu cầu phi chức năng", "15"),
        ("2.2. Đặc tả phần mềm", "16"),
        ("2.2.1. Phát biểu bài toán", "16"),
        ("2.2.2. Mô hình Use Case", "17"),
        ("2.2.3. Sơ đồ Sequence", "23"),
        ("CHƯƠNG 3. THIẾT KẾ CƠ SỞ DỮ LIỆU", "28"),
        ("3.1. Sơ đồ quan hệ giữa các bảng", "28"),
        ("3.2. Chi tiết các bảng", "29"),
        ("3.3. Kiến trúc hệ thống", "34"),
        ("3.3.1. Tổng quan kiến trúc hệ thống", "34"),
        ("3.3.2. Sơ đồ kiến trúc", "35"),
        ("3.3.3. Mô tả thành phần trong kiến trúc", "36"),
        ("3.3.4. Luồng xử lý yêu cầu", "38"),
        ("3.3.5. Ưu điểm của kiến trúc", "39"),
        ("3.3.6. Hạn chế và hướng phát triển", "40"),
        ("3.4. Thiết kế API", "41"),
        ("3.4.1. Xác thực", "41"),
        ("3.4.2. Chat APIs", "41"),
        ("3.4.3. Booking APIs", "42"),
        ("3.4.4. Booking Bill APIs", "42"),
        ("3.4.5. Order APIs", "43"),
        ("3.4.6. Menu APIs", "43"),
        ("3.4.7. Staff Request APIs", "43"),
        ("3.4.8. Staff Chat APIs", "44"),
        ("3.4.9. Realtime Notification APIs", "44"),
        ("3.4.10. Payment APIs", "45"),
        ("3.4.11. Mô hình bảo mật API", "46"),
        ("CHƯƠNG 4. HIỆN THỰC CHƯƠNG TRÌNH", "47"),
        ("4.1. Cấu trúc thư mục dự án", "47"),
        ("4.2. Hiện thực Backend", "48"),
        ("4.2.1. AI Agent Implementation", "48"),
        ("4.2.2. Knowledge Graph Pipeline", "50"),
        ("4.2.3. Database Models", "52"),
        ("4.2.4. API Endpoints", "53"),
        ("4.2.5. Payment Implementation", "54"),
        ("4.2.6. Docker Deployment", "56"),
        ("4.3. Hiện thực Frontend (Flutter App)", "57"),
        ("4.3.1. Cấu trúc Feature-based", "57"),
        ("4.3.2. State Management với Riverpod", "58"),
        ("4.3.3. Navigation với GoRouter", "58"),
        ("4.3.4. Chat UI với flutter_chat_ui", "59"),
        ("4.3.5. Giới thiệu giao diện ứng dụng", "59"),
        ("4.4. Môi trường triển khai", "65"),
        ("4.5. Kết quả đạt được", "66"),
        ("4.6. Hướng phát triển", "67"),
        ("TÀI LIỆU THAM KHẢO", "68"),
        ("PHỤ LỤC", "69"),
    ]

    for entry, page in toc_entries:
        p = doc.add_paragraph()
        if entry.startswith("CHƯƠNG") or entry in ["LỜI MỞ ĐẦU", "NHIỆM VỤ THIẾT KẾ", "MỤC LỤC", "DANH MỤC CÁC BẢNG BIỂU", "BẢNG TỪ VIẾT TẮT", "TÀI LIỆU THAM KHẢO", "PHỤ LỤC"]:
            run = p.add_run(entry)
            run.bold = True
        else:
            run = p.add_run(entry)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        set_paragraph_format(p, line_spacing=1.5)

    add_page_break(doc)

    # ====================================================================
    # DANH MỤC CÁC BẢNG BIỂU (không đánh số trang)
    # ====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DANH MỤC CÁC BẢNG BIỂU")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # Danh mục bảng
    add_body_text(doc, "Danh mục các bảng:", bold=True)

    bang_muc = [
        ("Bảng 1.1", "Môi trường phát triển"),
        ("Bảng 2.1", "Yêu cầu chức năng"),
        ("Bảng 2.2", "Yêu cầu phi chức năng"),
        ("Bảng 2.3", "Danh sách Actor"),
        ("Bảng 2.4", "Danh sách Use Case"),
        ("Bảng 3.1", "Bảng users"),
        ("Bảng 3.2", "Bảng bookings"),
        ("Bảng 3.3", "Bảng orders"),
        ("Bảng 3.4", "Bảng venues"),
        ("Bảng 3.5", "Bảng menu_items"),
        ("Bảng 3.6", "Bảng payments"),
        ("Bảng 3.7", "Bảng staff_requests"),
        ("Bảng 3.8", "Bảng notifications"),
        ("Bảng 3.9", "Authentication APIs"),
        ("Bảng 3.10", "Chat APIs"),
        ("Bảng 3.11", "Booking APIs"),
        ("Bảng 3.12", "Booking Bill APIs"),
        ("Bảng 3.13", "Order APIs"),
        ("Bảng 3.14", "Menu APIs"),
        ("Bảng 3.15", "Staff Request APIs"),
        ("Bảng 3.16", "Staff Chat APIs"),
        ("Bảng 3.17", "Realtime Notification APIs"),
        ("Bảng 3.18", "VNPay Payment APIs"),
        ("Bảng 3.19", "Stripe Payment APIs"),
        ("Bảng 4.1", "Môi trường triển khai"),
        ("Bảng 4.2", "Docker Services"),
        ("Bảng 4.3", "Flutter Dependencies"),
    ]

    table = doc.add_table(rows=len(bang_muc) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(["Mã bảng", "Tên bảng"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")

    for i, (ma, ten) in enumerate(bang_muc):
        table.rows[i + 1].cells[0].text = ma
        table.rows[i + 1].cells[1].text = ten
        for cell in table.rows[i + 1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    doc.add_paragraph()

    # Danh mục hình/sơ đồ
    add_body_text(doc, "Danh mục các sơ đồ, hình:", bold=True)

    hinh_muc = [
        ("Hình 1.1", "Sơ đồ kiến trúc tổng thể hệ thống"),
        ("Hình 2.1", "Use Case Diagram - Customer"),
        ("Hình 2.2", "Use Case Diagram - Staff"),
        ("Hình 2.3", "Use Case Diagram - Admin"),
        ("Hình 2.4", "Sequence Diagram - Chat"),
        ("Hình 2.5", "Sequence Diagram - Booking"),
        ("Hình 2.6", "Sequence Diagram - Payment (Stripe)"),
        ("Hình 2.7", "Sequence Diagram - Payment (VNPay)"),
        ("Hình 3.1", "Sơ đồ quan hệ các bảng (ERD)"),
        ("Hình 3.2", "Sơ đồ kiến trúc hệ thống"),
        ("Hình 3.3", "Luồng xử lý yêu cầu"),
        ("Hình 4.1", "Cấu trúc thư mục backend"),
        ("Hình 4.2", "Cấu trúc thư mục frontend"),
        ("Hình 4.3", "Luồng xử lý AI Agent"),
        ("Hình 4.4", "Luồng thanh toán Stripe"),
        ("Hình 4.5", "Luồng thanh toán VNPay"),
        ("Hình 4.6", "Kiến trúc Docker Deployment"),
        ("Hình 4.7", "Màn hình đăng nhập"),
        ("Hình 4.8", "Màn hình chính (Home Screen)"),
        ("Hình 4.9", "Giao diện Chat với AI Agent"),
        ("Hình 4.10", "Màn hình đặt sân (Booking)"),
        ("Hình 4.11", "Màn hình chọn thời gian đặt sân"),
        ("Hình 4.12", "Màn hình Menu đồ uống"),
        ("Hình 4.13", "Màn hình giỏ hàng (Cart)"),
        ("Hình 4.14", "Dialog chọn phương thức thanh toán"),
        ("Hình 4.15", "Màn hình thanh toán Stripe (Payment Sheet)"),
        ("Hình 4.16", "Màn hình thanh toán VNPay"),
        ("Hình 4.17", "Kết quả thanh toán thành công"),
        ("Hình 4.18", "Màn hình gọi nhân viên (Call Staff Dialog)"),
        ("Hình 4.19", "Màn hình danh sách yêu cầu (Staff Requests)"),
        ("Hình 4.20", "Màn hình quản lý yêu cầu (Staff Request Management)"),
        ("Hình 4.21", "Màn hình chat với nhân viên (Staff Chat)"),
        ("Hình 4.22", "Màn hình hộp thư nhân viên (Staff Inbox)"),
        ("Hình 4.23", "Màn hình thông báo vận hành (Notifications)"),
        ("Hình 4.24", "Màn hình hóa đơn đặt sân (Booking Bill)"),
        ("Hình 4.25", "Màn hình lịch sử đơn hàng (Order History)"),
        ("Hình 4.26", "Màn hình quản lý đặt sân (Admin Booking)"),
        ("Hình 4.27", "Màn hình quản lý menu (Admin Menu)"),
        ("Hình 4.28", "Màn hình Dashboard Admin"),
        ("Hình 4.29", "Màn hình hồ sơ cá nhân (Profile)"),
        ("Hình 4.30", "Thanh điều hướng nhân viên (Staff Shell)"),
    ]

    table2 = doc.add_table(rows=len(hinh_muc) + 1, cols=2)
    table2.style = "Table Grid"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["Mã hình", "Tên hình"]):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")

    for i, (ma, ten) in enumerate(hinh_muc):
        table2.rows[i + 1].cells[0].text = ma
        table2.rows[i + 1].cells[1].text = ten
        for cell in table2.rows[i + 1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    doc.add_paragraph()

    # ===== BẢNG TỪ VIẾT TẮT =====
    add_body_text(doc, "Bảng từ viết tắt:", bold=True)

    abbrev_data = [
        ("AI", "Artificial Intelligence - Trí tuệ nhân tạo"),
        ("API", "Application Programming Interface - Giao diện lập trình ứng dụng"),
        ("BWF", "Badminton World Federation - Liên đoàn Cầu lông Thế giới"),
        ("CRUD", "Create, Read, Update, Delete - Tạo, Đọc, Sửa, Xóa"),
        ("CSS", "Cascading Style Sheets"),
        ("DIO", "Dart HTTP client library"),
        ("ERD", "Entity-Relationship Diagram - Sơ đồ thực thể quan hệ"),
        ("FK", "Foreign Key - Khóa ngoại"),
        ("gRPC", "Google Remote Procedure Call"),
        ("HTTP", "Hypertext Transfer Protocol"),
        ("IDE", "Integrated Development Environment - Môi trường phát triển tích hợp"),
        ("JSON", "JavaScript Object Notation"),
        ("JWT", "JSON Web Token - Mã xác thực người dùng"),
        ("KG", "Knowledge Graph - Đồ thị tri thức"),
        ("LLM", "Large Language Model - Mô hình ngôn ngữ lớn"),
        ("ORM", "Object-Relational Mapping - Ánh xạ đối tượng-quan hệ"),
        ("PK", "Primary Key - Khóa chính"),
        ("RAG", "Retrieval Augmented Generation - Truy xuất tăng cường"),
        ("REST", "Representational State Transfer"),
        ("SDK", "Software Development Kit - Bộ phát triển phần mềm"),
        ("SQL", "Structured Query Language - Ngôn ngữ truy vấn có cấu trúc"),
        ("SQLAlchemy", "Python ORM library cho PostgreSQL"),
        ("UUID", "Universally Unique Identifier - Định danh duy nhất toàn cầu"),
        ("USAPA", "USA Pickleball Association - Hiệp hội Pickleball Hoa Kỳ"),
        ("VND", "Vietnamese Dong - Đồng Việt Nam"),
        ("VNPay", "Cổng thanh toán trực tuyến Việt Nam"),
        ("WebSocket", "Giao thức truyền thông hai chiều thời gian thực"),
        ("WPA", "World Pool-Billiard Association - Hiệp hội Bida Thế giới"),
    ]

    table_abbr = doc.add_table(rows=len(abbrev_data) + 1, cols=2)
    table_abbr.style = "Table Grid"
    table_abbr.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["Viết tắt", "Giải thích"]):
        cell = table_abbr.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")

    for i, (abbr, desc) in enumerate(abbrev_data):
        table_abbr.rows[i + 1].cells[0].text = abbr
        table_abbr.rows[i + 1].cells[1].text = desc
        for cell in table_abbr.rows[i + 1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    add_page_break(doc)

    # ====================================================================
    # CHƯƠNG 1. TỔNG QUAN (bắt đầu đánh số trang từ đây)
    # ====================================================================
    add_heading_custom(doc, "CHƯƠNG 1. TỔNG QUAN", level=1)

    # 1.1. Đặt vấn đề
    add_heading_custom(doc, "1.1. Đặt vấn đề", level=2)

    add_body_text(doc,
        "Trong những năm gần đây, ngành công nghiệp thể thao giải trí tại Việt Nam "
        "đang phát triển mạnh mẽ, đặc biệt là các môn thể thao như bida (billiards), "
        "pickleball và cầu lông. Số lượng quán thể thao, sân bãi tăng lên đáng kể, "
        "đi kèm với nhu cầu quản lý và phục vụ khách hàng ngày càng cao."
    )

    add_body_text(doc,
        "Tuy nhiên, hầu hết các quán thể thao hiện nay vẫn quản lý theo cách thủ công "
        "hoặc sử dụng các hệ thống đơn giản, thiếu tính tương tác và tự động hóa. "
        "Khách hàng thường gặp khó khăn trong việc tra cứu luật chơi, kỹ thuật, "
        "đặt sân, gọi đồ uống hay yêu cầu hỗ trợ từ nhân viên. Điều này ảnh hưởng "
        "trực tiếp đến trải nghiệm khách hàng và hiệu quả kinh doanh của quán."
    )

    add_body_text(doc,
        "Trước thực tế đó, việc xây dựng một hệ thống chatbot AI thông minh, "
        "có khả năng trả lời tự động các câu hỏi về luật chơi, kỹ thuật, "
        "hỗ trợ đặt sân, gọi đồ uống, thanh toán trực tuyến và quản lý yêu cầu khách hàng là cần thiết. "
        "Hệ thống không chỉ giúp tối ưu hóa quy trình phục vụ mà còn nâng cao "
        "trải nghiệm người dùng, tạo lợi thế cạnh tranh cho quán thể thao."
    )

    add_body_text(doc,
        "Đồ án này tập trung vào việc xây dựng Chatbot AI cho quán thể thao "
        "tích hợp Knowledge Graph và Tool Calling với các tính năng chính: "
        "tra cứu kiến thức thể thao (luật, kỹ thuật), đặt sân, gọi đồ uống, "
        "gọi nhân viên, thanh toán trực tuyến (Stripe, VNPay) và quản lý yêu cầu khách hàng."
    )

    # 1.2. Nhiệm vụ của đồ án
    add_heading_custom(doc, "1.2. Nhiệm vụ của đồ án", level=2)

    # 1.2.1. Mục đích
    add_heading_custom(doc, "1.2.1. Mục đích", level=3)

    add_body_text(doc,
        "Đồ án nhằm xây dựng một hệ thống chatbot AI ứng dụng cho quán thể thao "
        "(bida, pickleball, cầu lông) với các mục đích sau:"
    )

    purposes = [
        "Xây dựng chatbot AI có khả năng trả lời tự động các câu hỏi về luật chơi, kỹ thuật của các môn thể thao thông qua Knowledge Graph.",
        "Phát triển ứng dụng di động Flutter đa nền tảng (iOS, Android) với giao diện chat thân thiện, hỗ trợ đặt sân, gọi đồ uống và quản lý yêu cầu.",
        "Xây dựng backend API với FastAPI (Python) xử lý các nghiệp vụ: xác thực, đặt sân, đặt đồ uống, quản lý nhân viên và tích hợp AI Agent.",
        "Tích hợp công nghệ Tool Calling với LangChain để chatbot có thể tự động gọi các hàm xử lý nghiệp vụ khi phát hiện ý định người dùng.",
        "Sử dụng Knowledge Graph (Neo4j) để lưu trữ và truy vấn tri thức thể thao một cách có cấu trúc, hỗ trợ tìm kiếm ngữ nghĩa.",
        "Tích hợp thanh toán trực tuyến qua Stripe và VNPay, hỗ trợ đa phương thức thanh toán cho khách hàng.",
        "Triển khai hệ thống bằng Docker Compose với kiến trúc microservice.",
    ]
    for p in purposes:
        add_bullet_point(doc, p)

    # 1.2.2. Yêu cầu
    add_heading_custom(doc, "1.2.2. Yêu cầu", level=3)

    add_body_text(doc, "Yêu cầu chức năng:", bold=True)

    func_reqs = [
        "Chatbot trả lời câu hỏi về luật chơi, kỹ thuật của bida, pickleball, cầu lông.",
        "Đặt sân thể thao (billiards, pickleball, badminton) với xác nhận thời gian thực.",
        "Đặt đồ uống, đồ ăn từ menu của quán.",
        "Gọi nhân viên hỗ trợ khi cần thiết.",
        "Xem lịch sử đặt sân và trạng thái đơn hàng.",
        "Đăng nhập/đăng ký tài khoản khách hàng.",
        "Thanh toán trực tuyến qua Stripe hoặc VNPay.",
        "Quản lý danh sách sân, menu đồ uống cho admin.",
        "Quản lý giá dịch vụ và xem hóa đơn chi tiết cho admin.",
        "Thông báo real-time cho nhân viên khi có yêu cầu mới.",
    ]
    for r in func_reqs:
        add_bullet_point(doc, r)

    add_body_text(doc, "Yêu cầu phi chức năng:", bold=True)

    non_func_reqs = [
        "Ứng dụng hoạt động mượt mà trên cả iOS và Android.",
        "Thời gian phản hồi chatbot nhỏ hơn 5 giây.",
        "Bảo mật thông tin người dùng (mã hóa mật khẩu, JWT authentication).",
        "Bảo mật thanh toán (webhook signature verification, payment deduplication).",
        "Hệ thống có khả năng mở rộng (multi-tenant architecture).",
        "Giao diện thân thiện, dễ sử dụng.",
    ]
    for r in non_func_reqs:
        add_bullet_point(doc, r)

    # 1.2.3. Môi trường phát triển
    add_heading_custom(doc, "1.2.3. Môi trường phát triển", level=3)

    env_headers = ["Thành phần", "Công nghệ / Phiên bản"]
    env_data = [
        ["Hệ điều hành", "Windows 10/11, Ubuntu (WSL2)"],
        ["Ngôn ngữ lập trình", "Python 3.10+, Dart 3.1+, Java 17+"],
        ["Backend Framework", "FastAPI 0.115+"],
        ["Mobile Framework", "Flutter 3.1+"],
        ["AI Framework", "LangChain 0.3.1 + Ollama"],
        ["Cơ sở dữ liệu", "PostgreSQL 14+, Neo4j 5.x, Redis 7+"],
        ["Thanh toán", "Stripe SDK (Python), VNPay SDK (Java/Flutter)"],
        ["Triển khai", "Docker, Docker Compose"],
        ["IDE / Editor", "VS Code, Android Studio, IntelliJ IDEA"],
    ]
    create_table_with_header(doc, env_headers, env_data)

    # 1.3. Giới thiệu công nghệ
    add_heading_custom(doc, "1.3. Giới thiệu công nghệ", level=2)

    # 1.3.1. Flutter
    add_heading_custom(doc, "1.3.1. Flutter", level=3)

    add_body_text(doc,
        "Flutter là framework phát triển ứng dụng di động đa nền tảng (cross-platform) "
        "được phát triển bởi Google. Flutter sử dụng ngôn ngữ lập trình Dart và cho phép "
        "phát triển ứng dụng chạy trên cả iOS, Android, web và desktop từ một codebase duy nhất."
    )

    add_body_text(doc,
        "Trong đồ án này, Flutter được sử dụng để xây dựng ứng dụng di động với các tính năng: "
        "giao diện chat, quản lý đặt sân, menu đồ uống, xác thực người dùng. "
        "Các thư viện chính sử dụng bao gồm: Riverpod (quản lý trạng thái), "
        "GoRouter (điều hướng), Dio (gọi HTTP), flutter_chat_ui (giao diện chat)."
    )

    # 1.3.2. FastAPI
    add_heading_custom(doc, "1.3.2. FastAPI (Python)", level=3)

    add_body_text(doc,
        "FastAPI là framework web hiện đại, tốc độ cao được xây dựng trên Python 3.7+. "
        "FastAPI hỗ trợ tự động tạo tài liệu API (Swagger UI), xác thực dữ liệu "
        "tự động với Pydantic, và hỗ trợ async/await native."
    )

    add_body_text(doc,
        "Trong đồ án này, FastAPI đóng vai trò backend API server, xử lý tất cả "
        "các request từ frontend: xác thực người dùng (JWT), quản lý đặt sân, "
        "đơn hàng, menu, và tích hợp AI Agent để xử lý chat."
    )

    # 1.3.3. LangChain + Ollama
    add_heading_custom(doc, "1.3.3. LangChain + Ollama (AI Agent)", level=3)

    add_body_text(doc,
        "LangChain là framework phát triển ứng dụng AI dựa trên Large Language Model (LLM). "
        "LangChain hỗ trợ tích hợp công cụ (Tool Calling), cho phép AI tự động gọi "
        "các hàm xử lý nghiệp vụ khi phát hiện ý định người dùng."
    )

    add_body_text(doc,
        "Ollama là công cụ chạy LLM trên máy cục bộ (local), hỗ trợ các mô hình "
        "như qwen2.5-coder:7b, llama3.1:8b. Sử dụng Ollama giúp giảm chi phí "
        "API, bảo mật dữ liệu và không phụ thuộc vào kết nối internet."
    )

    add_body_text(doc,
        "Trong đồ án này, LangChain kết hợp với Ollama (hoặc Gemini API) để xử lý "
        "chat, phát hiện ý định (intent) và gọi các tool: query_knowledge (tra cứu KG), "
        "book_court (đặt sân), order_food (gọi món), call_staff (gọi nhân viên)."
    )

    # 1.3.4. Neo4j
    add_heading_custom(doc, "1.3.4. Neo4j (Knowledge Graph)", level=3)

    add_body_text(doc,
        "Neo4j là hệ quản trị cơ sở dữ liệu đồ thị (Graph Database) phổ biến nhất. "
        "Neo4j lưu trữ dữ liệu dưới dạng nodes (nút) và relationships (mối quan hệ), "
        "phù hợp để biểu diễn tri thức có cấu trúc phức tạp."
    )

    add_body_text(doc,
        "Trong đồ án này, Neo4j được sử dụng để lưu trữ Knowledge Graph chứa "
        "tri thức về luật chơi, kỹ thuật của các môn thể thao (bida, pickleball, cầu lông). "
        "Knowledge Graph hiện có 418 nodes (Rule, Technique, Equipment, Sport, Concept, GameType) "
        "và 441 relationships (DUNG_DE, LIEN_QUAN, LA_LOAI, THUOC, SU_DUNG, QUY_DINH)."
    )

    # 1.3.5. PostgreSQL
    add_heading_custom(doc, "1.3.5. PostgreSQL", level=3)

    add_body_text(doc,
        "PostgreSQL là hệ quản trị cơ sở dữ liệu quan hệ (RDBMS) mã nguồn mở, "
        "được sử dụng rộng rãi trong các ứng dụng production. PostgreSQL hỗ trợ "
        "kiểu dữ liệu JSON, full-text search và có hiệu suất cao."
    )

    add_body_text(doc,
        "Trong đồ án này, PostgreSQL lưu trữ dữ liệu nghiệp vụ: người dùng (users), "
        "đặt sân (bookings), đơn hàng (orders), menu đồ uống, và thông tin quán (venues)."
    )

    # 1.3.6. Redis
    add_heading_custom(doc, "1.3.6. Redis", level=3)

    add_body_text(doc,
        "Redis là hệ thống cache và message broker tốc độ cao, sử dụng cấu trúc dữ liệu "
        "trong bộ nhớ (in-memory). Redis hỗ trợ pub/sub pattern cho thông báo real-time."
    )

    add_body_text(doc,
        "Trong đồ án này, Redis được sử dụng để lưu trữ phiên chat (session), "
        "quản lý thông báo nhân viên (staff notification) và caching dữ liệu."
    )

    # 1.3.7. Stripe
    add_heading_custom(doc, "1.3.7. Stripe", level=3)

    add_body_text(doc,
        "Stripe là nền tảng thanh toán trực tuyến hàng đầu thế giới, hỗ trợ "
        "xử lý thanh toán qua thẻ tín dụng/ghi nợ (Visa, Mastercard, JCB) "
        "và các phương thức thanh toán quốc tế. Stripe cung cấp SDK cho "
        "nhiều nền tảng và API RESTful để tích hợp thanh toán vào ứng dụng."
    )

    add_body_text(doc,
        "Trong đồ án này, Stripe được tích hợp để xử lý thanh toán quốc tế "
        "cho booking và order. Backend sử dụng Stripe Python SDK để tạo "
        "Checkout Session, frontend hiển thị trang thanh toán Stripe qua WebView. "
        "Stripe webhook được sử dụng để xác nhận thanh toán từ server-to-server, "
        "đảm bảo tính toàn vẹn dữ liệu thanh toán."
    )

    # 1.3.8. VNPay
    add_heading_custom(doc, "1.3.8. VNPay", level=3)

    add_body_text(doc,
        "VNPay là cổng thanh toán trực tuyến phổ biến tại Việt Nam, hỗ trợ "
        "thanh toán qua thẻ ATM nội địa, Internet Banking, QR code và ví điện tử. "
        "VNPay cung cấp SDK native cho mobile và API cho web, phù hợp với "
        "người dùng Việt Nam."
    )

    add_body_text(doc,
        "Trong đồ án này, VNPay được tích hợp thông qua kiến trúc microservice: "
        "một Java service đóng vai trò proxy/gateway xử lý giao thức VNPay, "
        "giao tiếp với Python backend qua gRPC. Frontend sử dụng VNPay Native SDK "
        "qua Flutter Method Channel để mở trang thanh toán trực tiếp trên thiết bị."
    )

    # 1.3.9. Docker
    add_heading_custom(doc, "1.3.9. Docker", level=3)

    add_body_text(doc,
        "Docker là nền tảng containerization cho phép đóng gói ứng dụng "
        "và các phụ thuộc vào container độc lập. Docker Compose cho phép "
        "định nghĩa và quản lý nhiều container dịch vụ cùng lúc."
    )

    add_body_text(doc,
        "Trong đồ án này, Docker được sử dụng để triển khai hệ thống với "
        "4 dịch vụ: PostgreSQL (cơ sở dữ liệu), Redis (cache/message broker), "
        "Java Payment Service (VNPay gateway), và Python Backend (FastAPI). "
        "Tất cả giao tiếp qua mạng nội bộ (bridge network), chỉ Python Backend "
        "được expose ra ngoài qua port 8000."
    )

    add_page_break(doc)

    # ====================================================================
    # CHƯƠNG 2. PHÂN TÍCH
    # ====================================================================
    add_heading_custom(doc, "CHƯƠNG 2. PHÂN TÍCH", level=1)

    # 2.1. Xác định yêu cầu
    add_heading_custom(doc, "2.1. Xác định yêu cầu", level=2)

    # 2.1.1. Yêu cầu chức năng
    add_heading_custom(doc, "2.1.1. Yêu cầu chức năng", level=3)

    func_headers = ["STT", "Chức năng", "Mô tả"]
    func_data = [
        ["1", "Chat AI", "Người dùng gửi tin nhắn, chatbot trả lời tự động dựa trên Knowledge Graph và LLM."],
        ["2", "Tra cứu kiến thức", "Chatbot trả lời câu hỏi về luật chơi, kỹ thuật của bida, pickleball, cầu lông."],
        ["3", "Đặt sân", "Người dùng đặt sân thể thao (billiards, pickleball, badminton) theo thời gian."],
        ["4", "Đặt đồ uống", "Người dùng đặt đồ uống/đồ ăn từ menu của quán thông qua chatbot."],
        ["5", "Gọi nhân viên", "Người dùng gửi yêu cầu gọi nhân viên hỗ trợ."],
        ["6", "Xem lịch sử", "Xem lịch sử đặt sân, trạng thái đơn hàng."],
        ["7", "Đăng nhập/Đăng ký", "Xác thực người dùng bằng số điện thoại và mật khẩu."],
        ["8", "Quản lý (Admin)", "Quản lý sân, menu, nhân viên, thống kê doanh thu."],
        ["9", "Thanh toán trực tuyến", "Thanh toán booking và order qua Stripe (quốc tế) hoặc VNPay (trong nước)."],
        ["10", "Quản lý giá dịch vụ", "Admin cấu hình giá sân, xem hóa đơn chi tiết booking."],
    ]
    create_table_with_header(doc, func_headers, func_data)

    # 2.1.2. Yêu cầu phi chức năng
    add_heading_custom(doc, "2.1.2. Yêu cầu phi chức năng", level=3)

    nf_headers = ["Yêu cầu", "Mô tả"]
    nf_data = [
        ["Hiệu suất", "Thời gian phản hồi chatbot nhỏ hơn 5 giây cho hầu hết các truy vấn."],
        ["Bảo mật", "Mã hóa mật khẩu (bcrypt), JWT authentication, CORS protection, webhook signature verification."],
        ["Khả năng mở rộng", "Kiến trúc multi-tenant, hỗ trợ nhiều quán thể thao trên cùng hệ thống."],
        ["Đa nền tảng", "Ứng dụng chạy trên cả iOS và Android từ một codebase Flutter."],
        ["Tính sẵn sàng", "Fallback mechanism: khi LLM chính lỗi, tự chuyển sang LLM dự phòng."],
        ["Thanh toán", "Hỗ trợ đa phương thức thanh toán: Stripe (quốc tế), VNPay (trong nước), xác nhận qua webhook."],
    ]
    create_table_with_header(doc, nf_headers, nf_data)

    # 2.2. Đặc tả phần mềm
    add_heading_custom(doc, "2.2. Đặc tả phần mềm", level=2)

    # 2.2.1. Phát biểu bài toán
    add_heading_custom(doc, "2.2.1. Phát biểu bài toán", level=3)

    add_body_text(doc,
        "Xây dựng hệ thống chatbot AI cho quán thể thao (bida, pickleball, cầu lông) "
        "có khả năng:"
    )

    problem_items = [
        "Trả lời tự động các câu hỏi về luật chơi, kỹ thuật thông qua Knowledge Graph (Neo4j) với Graph RAG.",
        "Hỗ trợ đặt sân thể thao theo thời gian thực, kiểm tra tình trạng sân trống.",
        "Hỗ trợ đặt đồ uống/đồ ăn từ menu của quán.",
        "Tích hợp thanh toán trực tuyến qua Stripe (quốc tế) và VNPay (trong nước).",
        "Gửi yêu cầu gọi nhân viên hỗ trợ qua hệ thống thông báo real-time.",
        "Quản lý tài khoản người dùng với phân quyền: Customer, Staff, Admin.",
        "Cung cấp ứng dụng di động Flutter đa nền tảng với giao diện chat thân thiện.",
    ]
    for item in problem_items:
        add_bullet_point(doc, item)

    # 2.2.2. Mô hình Use Case
    add_heading_custom(doc, "2.2.2. Mô hình Use Case", level=3)

    add_body_text(doc, "Các actor trong hệ thống:", bold=True)

    actor_headers = ["Actor", "Mô tả"]
    actor_data = [
        ["Customer", "Khách hàng sử dụng ứng dụng để chat, đặt sân, đặt đồ uống."],
        ["Staff", "Nhân viên quán nhận thông báo yêu cầu từ khách hàng."],
        ["Admin", "Quản trị viên quản lý hệ thống, sân, menu, nhân viên."],
    ]
    create_table_with_header(doc, actor_headers, actor_data)

    doc.add_paragraph()
    add_body_text(doc, "Các Use Case chính:", bold=True)

    uc_headers = ["UC ID", "Use Case", "Actor"]
    uc_data = [
        ["UC01", "Gửi tin nhắn chat", "Customer"],
        ["UC02", "Tra cứu kiến thức thể thao", "Customer"],
        ["UC03", "Đặt sân thể thao", "Customer"],
        ["UC04", "Đặt đồ uống", "Customer"],
        ["UC05", "Gọi nhân viên", "Customer"],
        ["UC06", "Xem lịch sử đặt sân", "Customer"],
        ["UC07", "Đăng nhập/Đăng ký", "Customer"],
        ["UC08", "Nhận thông báo yêu cầu", "Staff"],
        ["UC09", "Quản lý sân và menu", "Admin"],
        ["UC10", "Xem thống kê doanh thu", "Admin"],
        ["UC11", "Thanh toán trực tuyến", "Customer"],
        ["UC12", "Quản lý giá dịch vụ", "Admin"],
        ["UC13", "Xem hóa đơn booking", "Admin"],
    ]
    create_table_with_header(doc, uc_headers, uc_data)

    doc.add_paragraph()
    add_body_text(doc, "Mô tả chi tiết Use Case UC01 - Gửi tin nhắn chat:", bold=True)

    uc_detail_headers = ["Thông tin", "Mô tả"]
    uc_detail_data = [
        ["Use Case", "UC01 - Gửi tin nhắn chat"],
        ["Actor", "Customer"],
        ["Mô tả", "Người dùng gửi tin nhắn văn bản đến chatbot và nhận phản hồi tự động."],
        ["Precondition", "Người dùng đã đăng nhập vào ứng dụng."],
        ["Main Flow", "1. Người dùng nhập tin nhắn vào ô chat.\n2. Hệ thống gửi tin nhắn đến backend API.\n3. AI Agent phân tích ý định (intent).\n4. Agent gọi tool phù hợp hoặc truy vấn Knowledge Graph.\n5. Hệ thống trả lời người dùng."],
        ["Alternative Flow", "3a. Nếu LLM chính lỗi, hệ thống fallback sang LLM dự phòng.\n4a. Nếu không xác định được ý định, chatbot trả lời mặc định."],
        ["Postcondition", "Tin nhắn được lưu vào session, người dùng nhận được phản hồi."],
    ]
    create_table_with_header(doc, uc_detail_headers, uc_detail_data)

    doc.add_paragraph()
    add_body_text(doc, "[Chèn hình Use Case Diagram tại đây]", bold=True)

    # ── Use Case Mermaid diagrams (per feature) ──

    add_mermaid_block(doc, "Use Case - Chat AI:", """graph LR
    Customer((Customer))
    AI((AI Agent))

    Customer -->|Gửi tin nhắn| UC01[UC01: Chat với AI]
    Customer -->|Hỏi luật/ky thuật| UC02[UC02: Tra cứu kiến thức]
    UC01 --> AI
    UC02 --> AI
    AI -->|Trả lời tự động| Customer
    AI -->|Dùng tool| KG[(Neo4j KG)]
""")

    add_mermaid_block(doc, "Use Case - Đặt sân:", """graph LR
    Customer((Customer))
    System((Hệ thống))

    Customer -->|Chọn loại sân| UC03[UC03: Đặt sân]
    Customer -->|Xem sân trống| UC03a[UC03a: Kiểm tra availability]
    Customer -->|Chọn thời gian| UC03b[UC03b: Chọn time slot]
    UC03 --> System
    UC03a --> System
    System -->|Xác nhận| Customer
    System -->|Tạo booking| DB[(PostgreSQL)]
""")

    add_mermaid_block(doc, "Use Case - Đặt đồ uống:", """graph LR
    Customer((Customer))
    System((Hệ thống))

    Customer -->|Xem menu| UC04[UC04: Đặt đồ uống]
    Customer -->|Chọn món| UC04a[UC04a: Thêm vào giỏ]
    Customer -->|Xác nhận đơn| UC04b[UC04b: Đặt hàng]
    UC04 --> System
    UC04b --> System
    System -->|Tạo order| DB[(PostgreSQL)]
    System -->|Thông báo| Staff((Staff))
""")

    add_mermaid_block(doc, "Use Case - Gọi nhân viên:", """graph LR
    Customer((Customer))
    Staff((Staff))
    System((Hệ thống))

    Customer -->|Chọn loại yêu cầu| UC05[UC05: Gọi nhân viên]
    Customer -->|Nhập mô tả| UC05a[UC05a: Tạo yêu cầu]
    UC05 --> System
    UC05a --> System
    System -->|Thông báo real-time| Staff
    Staff -->|Chấp nhận| UC08[UC08: Xử lý yêu cầu]
    Staff -->|Hoàn thành| UC08a[UC08a: Đóng yêu cầu]
""")

    add_mermaid_block(doc, "Use Case - Thanh toán:", """graph LR
    Customer((Customer))
    System((Hệ thống))
    Stripe((Stripe))
    VNPay((VNPay))

    Customer -->|Chọn thanh toán| UC11[UC11: Thanh toán]
    Customer -->|Chọn Stripe| UC11a[UC11a: Stripe Payment]
    Customer -->|Chọn VNPay| UC11b[UC11b: VNPay Payment]
    UC11a --> System
    UC11b --> System
    System -->|PaymentIntent| Stripe
    System -->|Payment URL| VNPay
    Stripe -->|Webhook xác nhận| System
    VNPay -->|Callback| System
    System -->|Cập nhật trạng thái| Customer
""")

    add_mermaid_block(doc, "Use Case - Staff Chat:", """graph LR
    Customer((Customer))
    Staff((Staff))
    System((Hệ thống))

    Customer -->|Gửi tin nhắn| UC_C1[Chat với nhân viên]
    Staff -->|Gửi tin nhắn| UC_S1[Chat với khách hàng]
    UC_C1 --> System
    UC_S1 --> System
    System -->|WebSocket| Customer
    System -->|WebSocket| Staff
    Staff -->|Xem inbox| UC_S2[UC: Staff Inbox]
    Staff -->|Đóng phòng chat| UC_S3[UC: Đóng room]
""")

    add_mermaid_block(doc, "Use Case - Admin:", """graph LR
    Admin((Admin))

    Admin -->|Quản lý| UC09[UC09: Quản lý sân & menu]
    Admin -->|Xem thống kê| UC10[UC10: Dashboard]
    Admin -->|Quản lý đặt sân| UC12[UC12: Quản lý booking]
    Admin -->|Xem hóa đơn| UC13[UC13: Xem bill]
    Admin -->|Quản lý đơn hàng| UC14[UC14: Quản lý order]
    Admin -->|CRUD menu| UC15[UC15: Quản lý menu]
""")

    doc.add_paragraph()
    add_body_text(doc, "Mô tả chi tiết Use Case UC11 - Thanh toán trực tuyến:", bold=True)

    uc_payment_headers = ["Thông tin", "Mô tả"]
    uc_payment_data = [
        ["Use Case", "UC11 - Thanh toán trực tuyến"],
        ["Actor", "Customer"],
        ["Mô tả", "Người dùng thanh toán booking hoặc order qua Stripe hoặc VNPay."],
        ["Precondition", "Người dùng đã tạo booking/order thành công, chưa thanh toán."],
        ["Main Flow", "1. Người dùng chọn 'Thanh toán ngay' sau khi đặt sân/đồ uống.\n2. Hệ thống hiển thị dialog chọn phương thức: Stripe hoặc VNPay.\n3a. Nếu Stripe: Hệ thống tạo Checkout Session, mở WebView thanh toán.\n3b. Nếu VNPay: Hệ thống tạo payment URL, mở VNPay Native SDK.\n4. Người dùng hoàn tất thanh toán trên cổng thanh toán.\n5. Hệ thống nhận xác nhận (webhook/callback), cập nhật trạng thái."],
        ["Alternative Flow", "3a. Nếu Stripe lỗi, hiển thị thông báo và cho phép thử lại.\n4a. Nếu người dùng hủy thanh toán, trạng thái giữ nguyên 'unpaid'."],
        ["Postcondition", "Trạng thái thanh toán được cập nhật: 'paid' hoặc 'failed'."],
    ]
    create_table_with_header(doc, uc_payment_headers, uc_payment_data)

    # 2.2.3. Sơ đồ Sequence
    add_heading_custom(doc, "2.2.3. Sơ đồ Sequence", level=3)

    add_body_text(doc,
        "Luồng xử lý khi người dùng gửi tin nhắn chat (Sequence Diagram):"
    )

    seq_items = [
        "1. Người dùng (Customer) nhập tin nhắn vào Flutter App.",
        "2. Flutter App gửi POST request đến /api/chat với message và session_id.",
        "3. ChatService nhận request và gọi VenueAgent.process().",
        "4. IntentRouter phân tích ý định người dùng (embedding-based classification).",
        "5. Nếu ý định là câu hỏi đơn giản (greeting, goodbye), IntentRouter trả lời trực tiếp.",
        "6. Nếu ý định phức tạp, VenueAgent gọi LLM (Ollama/Gemini) với system prompt.",
        "7. LLM phân tích và quyết định gọi tool phù hợp (query_knowledge, book_court, order_food, call_staff).",
        "8. Tool được thực thi: query_knowledge truy vấn Neo4j, book_court insert PostgreSQL.",
        "9. Kết quả tool được trả về LLM để tổng hợp phản hồi.",
        "10. Phản hồi được gửi trả về Flutter App qua HTTP response.",
    ]
    for item in seq_items:
        add_body_text(doc, item)

    doc.add_paragraph()
    add_body_text(doc, "Luồng xử lý đặt sân (Booking Sequence):")

    booking_seq = [
        "1. Người dùng yêu cầu đặt sân qua chat: 'Đặt sân bida lúc 7h tối mai'.",
        "2. AI Agent phát hiện intent booking, gọi tool book_court.",
        "3. Tool book_court kiểm tra sân trống trong database PostgreSQL.",
        "4. Nếu sân trống, tạo bản ghi booking mới với status 'confirmed'.",
        "5. Trả về xác nhận đặt sân cho người dùng qua chatbot.",
        "6. Hệ thống hiển thị dialog thanh toán với 2 lựa chọn: Stripe hoặc VNPay.",
    ]
    for item in booking_seq:
        add_body_text(doc, item)

    doc.add_paragraph()
    add_body_text(doc, "Luồng xử lý thanh toán (Payment Sequence - Stripe):")

    payment_seq = [
        "1. Người dùng chọn 'Thanh toán ngay' và chọn phương thức Stripe.",
        "2. Frontend gọi POST /api/stripe/create-checkout với order_id, amount.",
        "3. Backend tạo Stripe Checkout Session, trả về checkout_url.",
        "4. Frontend mở WebView hiển thị trang thanh toán Stripe.",
        "5. Người dùng nhập thông tin thẻ và hoàn tất thanh toán.",
        "6. Stripe gửi webhook POST /api/stripe/webhook xác nhận.",
        "7. Backend xác thực chữ ký, cập nhật payment_status = 'paid'.",
        "8. Frontend redirect về trang kết quả thanh toán.",
    ]
    for item in payment_seq:
        add_body_text(doc, item)

    doc.add_paragraph()
    add_body_text(doc, "[Chèn hình Sequence Diagram tại đây]", bold=True)

    # ── Sequence Mermaid diagrams (per feature) ──

    add_mermaid_block(doc, "Sequence - Chat với AI Agent:", """sequenceDiagram
    participant C as Customer
    participant F as Flutter App
    participant B as FastAPI Backend
    participant A as AI Agent
    participant KG as Neo4j KG
    participant R as Redis

    C->>F: Nhập tin nhắn
    F->>B: POST /api/chat {message, session_id}
    B->>R: Lấy session history
    R-->>B: History
    B->>A: process(message, history)
    A->>A: IntentRouter phân tích intent
    alt Intent đơn giản (greeting)
        A-->>B: Trả lời trực tiếp
    else Intent phức tạp
        A->>KG: query_knowledge (Cypher)
        KG-->>A: Kết quả
        A->>A: LLM tổng hợp phản hồi
        A-->>B: Response
    end
    B-->>F: ChatResponse {reply, tools_used}
    F-->>C: Hiển thị tin nhắn
    B->>R: Lưu session
""")

    add_mermaid_block(doc, "Sequence - Đặt sân (Booking):", """sequenceDiagram
    participant C as Customer
    participant F as Flutter App
    participant B as FastAPI Backend
    participant DB as PostgreSQL

    C->>F: Chọn loại sân + thời gian
    F->>B: GET /api/booking/availability
    B->>DB: Query sân trống
    DB-->>B: Danh sách sân
    B-->>F: AvailabilityResponse
    F-->>C: Hiển thị sân trống

    C->>F: Xác nhận đặt sân
    F->>B: POST /api/booking/
    B->>DB: Tạo booking (status=confirmed)
    DB-->>B: Booking created
    B-->>F: BookingResponse
    F-->>C: Đặt sân thành công
    F->>F: Hiển thị dialog thanh toán
""")

    add_mermaid_block(doc, "Sequence - Đặt đồ uống (Order):", """sequenceDiagram
    participant C as Customer
    participant F as Flutter App
    participant B as FastAPI Backend
    participant DB as PostgreSQL
    participant S as Staff

    C->>F: Chọn món từ menu
    F->>F: Thêm vào giỏ hàng
    C->>F: Xác nhận đặt hàng
    F->>B: POST /api/order/ {items, booking_id}
    B->>DB: Tạo order + order_items
    B->>DB: Link order với booking (nếu có)
    DB-->>B: Order created
    B-->>F: OrderResponse
    F-->>C: Đặt hàng thành công
    B->>S: Thông báo đơn hàng mới (WebSocket)
""")

    add_mermaid_block(doc, "Sequence - Thanh toán Stripe (PaymentIntent):", """sequenceDiagram
    participant C as Customer
    participant F as Flutter App
    participant FS as flutter_stripe SDK
    participant B as FastAPI Backend
    participant S as Stripe API
    participant DB as PostgreSQL

    C->>F: Chọn thanh toán Stripe
    F->>B: POST /api/stripe/create-payment-intent {order_type, order_id, amount}
    B->>S: stripe.PaymentIntent.create(amount, currency)
    S-->>B: PaymentIntent {client_secret}
    B-->>F: {client_secret, payment_intent_id}
    F->>FS: initPaymentSheet(clientSecret)
    FS->>C: Hiển thị Payment Sheet (native)
    C->>FS: Nhập thông tin thẻ, xác nhận
    FS->>S: Xác nhận thanh toán
    S-->>FS: Payment thành công
    FS-->>F: PaymentSheetResponse
    F->>B: POST /api/stripe/webhook (payment_intent.succeeded)
    B->>DB: Cập nhật payment_status = 'paid'
    B-->>F: Redirect /success
    F-->>C: Thanh toán thành công
""")

    add_mermaid_block(doc, "Sequence - Thanh toán VNPay:", """sequenceDiagram
    participant C as Customer
    participant F as Flutter App
    participant B as FastAPI Backend
    participant JV as Java Payment Service
    participant V as VNPay API

    C->>F: Chọn thanh toán VNPay
    F->>B: POST /api/payment/create {order_type, order_id, amount}
    B->>JV: gRPC: createPayment
    JV->>V: Tạo VNPay payment URL
    V-->>JV: payment_url
    JV-->>B: payment_url
    B-->>F: {payment_url}
    F->>F: Mở VNPay Native SDK (MethodChannel)
    C->>V: Chọn ngân hàng, hoàn tất
    V-->>B: Callback URL
    B->>B: Xác thực chữ ký VNPay
    B->>DB: Cập nhật payment_status = 'paid'
    B-->>F: Kết quả thanh toán
    F-->>C: Thanh toán thành công
""")

    add_mermaid_block(doc, "Sequence - Gọi nhân viên (Staff Request):", """sequenceDiagram
    participant C as Customer
    participant F as Flutter App
    participant B as FastAPI Backend
    participant DB as PostgreSQL
    participant WS as WebSocket
    participant S as Staff

    C->>F: Mở dialog gọi nhân viên
    C->>F: Chọn loại yêu cầu + mô tả
    F->>B: POST /api/staff/requests
    B->>DB: Tạo staff_request
    B->>WS: broadcast_to_roles [STAFF]
    WS-->>S: Thông báo yêu cầu mới
    S->>F: Xem danh sách yêu cầu
    S->>B: PATCH /requests/{id}/accept
    B->>DB: Cập nhật status=accepted
    B->>WS: Thông báo cho customer
    WS-->>C: Yêu cầu đã được chấp nhận
    S->>B: PATCH /requests/{id}/complete
    B->>DB: Cập nhật status=completed
    B->>WS: Thông báo hoàn thành
    WS-->>C: Yêu cầu đã hoàn thành
""")

    add_mermaid_block(doc, "Sequence - Staff Chat (Real-time):", """sequenceDiagram
    participant C as Customer
    participant CW as Customer WebSocket
    participant B as FastAPI Backend
    participant R as Redis
    participant SW as Staff WebSocket
    participant S as Staff

    Note over C,S: Staff request accepted -> Chat room created

    C->>CW: Kết nối /api/staff/chat/{room_id}/ws
    S->>SW: Kết nối /api/staff/chat/{room_id}/ws
    B->>R: Lưu room info

    C->>CW: {type: "message", content: "Xin chào"}
    CW->>B: WebSocket message
    B->>R: Lưu message (rpush)
    B->>SW: broadcast_to_room
    SW-->>S: Hiển thị tin nhắn

    S->>SW: {type: "message", content: "Chào bạn"}
    SW->>B: WebSocket message
    B->>R: Lưu message (rpush)
    B->>CW: broadcast_to_room
    CW-->>C: Hiển thị tin nhắn

    Note over C,S: Presence tracking qua Redis (TTL 60s)
    C->>CW: {type: "ping"}
    CW->>B: refresh_presence
    B->>R: SET online:{room}:{user} TTL=60s
""")

    add_mermaid_block(doc, "Sequence - Realtime Notification:", """sequenceDiagram
    participant F as Flutter App
    participant WS as WebSocket Client
    participant B as FastAPI Backend
    participant R as Redis
    participant N as NotificationService
    participant DB as PostgreSQL

    Note over F,DB: App khởi động -> Kết nối WebSocket

    F->>WS: Kết nối /api/realtime/notifications?token=xxx
    B->>B: Xác thực JWT token
    B->>B: connect(websocket, role, user_id)

    Note over F,DB: Khi có sự kiện mới (booking/order/request)

    B->>N: notify_operations(event_type, title, message)
    N->>DB: Lưu notification
    N->>R: SET + PUBLISH
    N->>B: broadcast_to_roles(target_roles)
    B->>WS: Gửi notification qua WebSocket
    WS-->>F: Nhận message
    F->>F: showOperationNotification (native)
    F->>F: Cập nhật danh sách in-app
""")

    add_page_break(doc)

    # ====================================================================
    # CHƯƠNG 3. THIẾT KẾ CƠ SỞ DỮ LIỆU
    # ====================================================================
    add_heading_custom(doc, "CHƯƠNG 3. THIẾT KẾ CƠ SỞ DỮ LIỆU", level=1)

    # 3.1. Sơ đồ quan hệ giữa các bảng
    add_heading_custom(doc, "3.1. Sơ đồ quan hệ giữa các bảng", level=2)

    add_body_text(doc,
        "Hệ thống sử dụng hai loại cơ sở dữ liệu chính:"
    )

    add_body_text(doc,
        "1. PostgreSQL (RDBMS): Lưu trữ dữ liệu nghiệp vụ có cấu trúc "
        "(users, bookings, orders, venues, menu)."
    )

    add_body_text(doc,
        "2. Neo4j (Graph DB): Lưu trữ Knowledge Graph tri thức thể thao "
        "(entities, relationships)."
    )

    doc.add_paragraph()
    add_body_text(doc, "Sơ đồ quan hệ các bảng trong PostgreSQL:", bold=True)

    add_body_text(doc,
        "users (1) --- (N) bookings\n"
        "users (1) --- (N) orders\n"
        "venues (1) --- (N) bookings\n"
        "venues (1) --- (N) service_resources\n"
        "orders (1) --- (N) order_items\n"
        "businesses (1) --- (N) venues\n"
        "businesses (1) --- (N) users\n"
        "bookings (1) --- (N) payments\n"
        "orders (1) --- (N) payments"
    )

    doc.add_paragraph()
    add_body_text(doc, "[Chèn hình ERD Diagram tại đây]", bold=True)

    # ── Class Diagram (Mermaid) ──
    add_mermaid_block(doc, "Sơ đồ Class - Backend Models:", """classDiagram
    class User {
        +UUID id
        +String phone
        +String name
        +String email
        +String password_hash
        +UserRole role
        +String stripe_customer_id
        +DateTime created_at
    }

    class Booking {
        +UUID id
        +String user_id
        +UUID venue_id
        +CourtType court_type
        +Int court_number
        +DateTime start_time
        +DateTime end_time
        +BookingStatus status
        +Numeric total_price
        +String payment_status
    }

    class Order {
        +UUID id
        +String user_id
        +UUID venue_id
        +UUID booking_id
        +Numeric total_price
        +OrderStatus status
        +String payment_status
        +DateTime created_at
    }

    class OrderItem {
        +UUID id
        +UUID order_id
        +UUID menu_item_id
        +String name
        +Int price
        +Int quantity
    }

    class MenuItem {
        +UUID id
        +UUID venue_id
        +String name
        +Int price
        +String category
        +Boolean available
    }

    class Venue {
        +UUID id
        +String name
        +String address
        +String phone
        +UUID business_id
    }

    class ServiceResource {
        +UUID id
        +UUID venue_id
        +String name
        +CourtType resource_type
        +Numeric hourly_rate
        +Boolean active
    }

    class Payment {
        +UUID id
        +String order_type
        +UUID order_id
        +Numeric amount
        +String bank_code
        +String vnp_transaction_no
        +String stripe_payment_intent_id
        +String status
        +DateTime paid_at
    }

    class StaffRequest {
        +UUID id
        +String user_id
        +UUID venue_id
        +UUID resource_id
        +RequestType request_type
        +String description
        +RequestStatus status
        +String accepted_by
        +DateTime created_at
        +DateTime accepted_at
        +DateTime completed_at
    }

    class Notification {
        +UUID id
        +String event_type
        +String title
        +String message
        +JSON target_roles
        +String source
        +JSON payload
        +DateTime read_at
        +DateTime created_at
    }

    User "1" -- "*" Booking
    User "1" -- "*" Order
    User "1" -- "*" StaffRequest
    Venue "1" -- "*" Booking
    Venue "1" -- "*" Order
    Venue "1" -- "*" MenuItem
    Venue "1" -- "*" ServiceResource
    Venue "1" -- "*" StaffRequest
    Order "1" -- "*" OrderItem
    Booking "1" -- "*" Payment
    Order "1" -- "*" Payment
""")

    # ── Class Diagram - Flutter Architecture ──
    add_mermaid_block(doc, "Sơ đồ Class - Flutter Feature Architecture:", """classDiagram
    class AuthProvider {
        +AsyncValue~User?~ state
        +login(phone, password)
        +logout()
        +refreshProfile()
    }

    class ChatProvider {
        +List~ChatMessage~ messages
        +sendMessage(text)
        +streamResponse()
    }

    class BookingProvider {
        +List~Booking~ bookings
        +createBooking()
        +checkAvailability()
        +getActiveBooking()
    }

    class CartNotifier {
        +List~CartItem~ items
        +addItem(item)
        +removeItem(id)
        +clear()
        +placeOrder()
    }

    class StripeNotifier {
        +initPaymentSheet()
        +pay(orderType, orderId, amount)
    }

    class StaffNotificationsNotifier {
        +List~StaffNotification~ notifications
        +WebSocketChannel channel
        +start()
        +stop()
        +markAsRead()
        +markAllAsRead()
    }

    class StaffRequestNotifier {
        +List~StaffRequest~ requests
        +createRequest()
        +acceptRequest()
        +completeRequest()
        +cancelRequest()
    }

    class CustomerChatNotificationsNotifier {
        +WebSocketChannel channel
        +start()
        +stop()
    }

    class LocalNotificationService {
        +MethodChannel channel
        +showOperationNotification(title, body)
    }

    AuthProvider --> ChatProvider : provides user
    AuthProvider --> BookingProvider : provides user
    BookingProvider --> CartNotifier : active booking
    CartNotifier --> StripeNotifier : triggers payment
    StaffNotificationsNotifier --> LocalNotificationService : shows native notif
    CustomerChatNotificationsNotifier --> LocalNotificationService : shows native notif
""")

    # 3.2. Chi tiết các bảng
    add_heading_custom(doc, "3.2. Chi tiết các bảng", level=2)

    # Bảng users
    add_body_text(doc, "Bảng users:", bold=True)

    users_headers = ["Tên cột", "Kiểu dữ liệu", "Khóa", "Mô tả"]
    users_data = [
        ["id", "UUID", "PK", "Khóa chính, tự sinh"],
        ["phone", "VARCHAR(20)", "UNIQUE, NOT NULL", "Số điện thoại (dùng để đăng nhập)"],
        ["name", "VARCHAR(100)", "NOT NULL", "Tên người dùng"],
        ["email", "VARCHAR(255)", "UNIQUE", "Email (tùy chọn)"],
        ["password_hash", "VARCHAR(255)", "", "Mật khẩu đã mã hóa (bcrypt)"],
        ["role", "ENUM", "NOT NULL", "CUSTOMER / STAFF / ADMIN"],
        ["created_at", "TIMESTAMP", "NOT NULL", "Thời gian tạo"],
    ]
    create_table_with_header(doc, users_headers, users_data)

    doc.add_paragraph()

    # Bảng bookings
    add_body_text(doc, "Bảng bookings:", bold=True)

    bookings_headers = ["Tên cột", "Kiểu dữ liệu", "Khóa", "Mô tả"]
    bookings_data = [
        ["id", "UUID", "PK", "Khóa chính"],
        ["user_id", "VARCHAR(128)", "FK -> users", "Người đặt sân"],
        ["venue_id", "UUID", "FK -> venues", "Quán thể thao"],
        ["court_type", "ENUM", "NOT NULL", "billiards / pickleball / badminton"],
        ["court_number", "INTEGER", "NOT NULL", "Số sân"],
        ["start_time", "TIMESTAMP", "NOT NULL", "Thời gian bắt đầu"],
        ["end_time", "TIMESTAMP", "NOT NULL", "Thời gian kết thúc"],
        ["status", "ENUM", "NOT NULL", "confirmed / cancelled / completed"],
        ["total_price", "NUMERIC(12,2)", "", "Tổng tiền (VND)"],
        ["payment_status", "VARCHAR(20)", "DEFAULT 'unpaid'", "Trạng thái thanh toán"],
    ]
    create_table_with_header(doc, bookings_headers, bookings_data)

    doc.add_paragraph()

    # Bảng orders
    add_body_text(doc, "Bảng orders:", bold=True)

    orders_headers = ["Tên cột", "Kiểu dữ liệu", "Khóa", "Mô tả"]
    orders_data = [
        ["id", "UUID", "PK", "Khóa chính"],
        ["user_id", "VARCHAR(128)", "FK -> users", "Người đặt hàng"],
        ["venue_id", "UUID", "FK -> venues", "Quán thể thao"],
        ["total_price", "NUMERIC(12,2)", "DEFAULT 0", "Tổng tiền (VND)"],
        ["status", "ENUM", "NOT NULL", "pending/preparing/ready/delivered/cancelled"],
        ["payment_status", "VARCHAR(20)", "DEFAULT 'unpaid'", "Trạng thái thanh toán"],
        ["created_at", "TIMESTAMP", "NOT NULL", "Thời gian đặt hàng"],
    ]
    create_table_with_header(doc, orders_headers, orders_data)

    doc.add_paragraph()

    # Bảng venues
    add_body_text(doc, "Bảng venues:", bold=True)

    venues_headers = ["Tên cột", "Kiểu dữ liệu", "Khóa", "Mô tả"]
    venues_data = [
        ["id", "UUID", "PK", "Khóa chính"],
        ["name", "VARCHAR(255)", "NOT NULL", "Tên quán thể thao"],
        ["address", "TEXT", "", "Địa chỉ"],
        ["phone", "VARCHAR(20)", "", "Số điện thoại liên hệ"],
        ["business_id", "UUID", "FK -> businesses", "Doanh nghiệp sở hữu"],
    ]
    create_table_with_header(doc, venues_headers, venues_data)

    doc.add_paragraph()

    # Bảng menu_items
    add_body_text(doc, "Bảng menu_items:", bold=True)

    menu_headers = ["Tên cột", "Kiểu dữ liệu", "Khóa", "Mô tả"]
    menu_data = [
        ["id", "UUID", "PK", "Khóa chính"],
        ["venue_id", "UUID", "FK -> venues", "Quán thể thao"],
        ["name", "VARCHAR(255)", "NOT NULL", "Tên món"],
        ["price", "INTEGER", "NOT NULL", "Giá (VND)"],
        ["category", "VARCHAR(100)", "", "Loại: đồ uống/đồ ăn"],
        ["available", "BOOLEAN", "NOT NULL", "Còn phục vụ hay không"],
    ]
    create_table_with_header(doc, menu_headers, menu_data)

    doc.add_paragraph()

    # Bảng payments
    add_body_text(doc, "Bảng payments:", bold=True)

    payments_headers = ["Tên cột", "Kiểu dữ liệu", "Khóa", "Mô tả"]
    payments_data = [
        ["id", "UUID", "PK", "Khóa chính"],
        ["order_type", "VARCHAR(20)", "NOT NULL", "Loại đối tượng: 'booking' hoặc 'order'"],
        ["order_id", "UUID", "NOT NULL", "ID của booking hoặc order"],
        ["amount", "NUMERIC(12,2)", "NOT NULL", "Số tiền thanh toán (VND)"],
        ["description", "VARCHAR(500)", "", "Mô tả thanh toán"],
        ["bank_code", "VARCHAR(50)", "", "Mã ngân hàng/cổng thanh toán (stripe, vnpay)"],
        ["vnp_transaction_no", "VARCHAR(100)", "", "Mã giao dịch VNPay"],
        ["status", "VARCHAR(20)", "NOT NULL", "pending/completed/failed"],
        ["paid_at", "TIMESTAMP", "", "Thời gian thanh toán thành công"],
        ["created_at", "TIMESTAMP", "NOT NULL", "Thời gian tạo bản ghi"],
    ]
    create_table_with_header(doc, payments_headers, payments_data)

    # 3.3. Kiến trúc hệ thống
    add_heading_custom(doc, "3.3. Kiến trúc hệ thống", level=2)

    # 3.3.1. Tổng quan kiến trúc hệ thống
    add_heading_custom(doc, "3.3.1. Tổng quan kiến trúc hệ thống", level=3)

    add_body_text(doc,
        "Hệ thống được thiết kế theo kiến trúc phân tầng (Layered Architecture) "
        "với các thành phần chính:"
    )

    arch_items = [
        "Presentation Layer: Flutter App (Mobile UI, Chat Interface)",
        "Application Layer: FastAPI Backend (REST API, Business Logic)",
        "Domain Layer: AI Agent (LangChain, Intent Router, Tool Calling)",
        "Infrastructure Layer: PostgreSQL, Neo4j, Redis",
    ]
    for item in arch_items:
        add_bullet_point(doc, item)

    # 3.3.2. Sơ đồ kiến trúc
    add_heading_custom(doc, "3.3.2. Sơ đồ kiến trúc", level=3)

    add_body_text(doc,
        "Sơ đồ kiến trúc tổng thể của hệ thống được trình bày trong hình bên dưới."
    )

    add_body_text(doc,
        "┌─────────────────────────────────────┐\n"
        "│           Flutter App               │  Presentation Layer\n"
        "│  (Chat UI, Booking, Menu, Auth)     │\n"
        "│  (Staff Chat, Staff Requests)       │\n"
        "│  (Payment: Stripe/VNPay SDK)        │\n"
        "└──────────────┬──────────────────────┘\n"
        "               │ HTTP / WebSocket\n"
        "┌──────────────▼──────────────────────┐\n"
        "│          FastAPI Backend            │  Application Layer\n"
        "│  /chat /booking /order /staff       │\n"
        "│  /payment /stripe /realtime         │\n"
        "│  /staff/chat /staff/requests        │\n"
        "└──────┬───────────────────┬──────────┘\n"
        "       │                   │\n"
        "┌──────▼──────┐    ┌───────▼─────────┐\n"
        "│  AI Agent   │    │ Payment Service │  Domain Layer\n"
        "│  LangChain  │    │  (Java/VNPay)   │\n"
        "│  Tool Call  │    │  Stripe SDK     │\n"
        "└──────┬──────┘    └───────┬─────────┘\n"
        "       │                   │\n"
        "┌──────▼───────────────────▼─────────┐\n"
        "│         Infrastructure Layer        │\n"
        "│  PostgreSQL │ Neo4j │ Redis │ Docker│\n"
        "└─────────────────────────────────────┘"
    )

    doc.add_paragraph()
    add_body_text(doc, "[Chèn hình Architecture Diagram tại đây]", bold=True)

    # ── System Architecture Flow (Mermaid) ──
    add_mermaid_block(doc, "Sơ đồ luồng tổng thể hệ thống:", """flowchart TB
    subgraph Mobile["Flutter App"]
        Chat[Chat UI]
        Booking[Booking Screen]
        Menu[Menu Screen]
        Payment[Payment Screen]
        StaffReq[Staff Request]
        StaffChat[Staff Chat]
        Notif[Notifications]
    end

    subgraph Backend["FastAPI Backend"]
        Auth[Auth API]
        ChatAPI[Chat API]
        BookAPI[Booking API]
        OrderAPI[Order API]
        MenuAPI[Menu API]
        StaffAPI[Staff Request API]
        ChatWS[Staff Chat WS]
        RealtimeWS[Realtime WS]
        PayAPI[Payment API]
        StripeAPI[Stripe API]
        AdminAPI[Admin API]
    end

    subgraph Domain["Domain Layer"]
        Agent[AI Agent\nLangChain + Tools]
        NotifSvc[Notification Service]
        StaffChatSvc[Staff Chat Service]
    end

    subgraph Infra["Infrastructure"]
        PG[(PostgreSQL)]
        Neo[(Neo4j KG)]
        Redis[(Redis)]
        Stripe[Stripe API]
        VNPay[VNPay API]
    end

    Mobile -->|HTTP/WS| Backend
    Backend --> Domain
    Domain --> Infra
    ChatAPI --> Agent
    Agent --> Neo
    BookAPI --> PG
    OrderAPI --> PG
    RealtimeWS --> NotifSvc
    ChatWS --> StaffChatSvc
    StaffChatSvc --> Redis
    NotifSvc --> PG
    StripeAPI --> Stripe
    PayAPI --> VNPay
""")

    # 3.3.3. Mô tả thành phần trong kiến trúc
    add_heading_custom(doc, "3.3.3. Mô tả thành phần trong kiến trúc", level=3)

    add_body_text(doc, "Flutter App (Presentation Layer):", bold=True)

    flutter_items = [
        "Giao diện chat: Sử dụng flutter_chat_ui để hiển thị tin nhắn, hỗ trợ markdown.",
        "Booking Screen: Chọn loại sân, thời gian, xem sân trống.",
        "Menu Screen: Hiển thị danh sách đồ uống, đặt hàng.",
        "Auth Screen: Đăng nhập/đăng ký với số điện thoại.",
        "State Management: Sử dụng Riverpod để quản lý trạng thái ứng dụng.",
    ]
    for item in flutter_items:
        add_bullet_point(doc, item)

    add_body_text(doc, "FastAPI Backend (Application Layer):", bold=True)

    backend_items = [
        "API Routers: chat, booking, order, menu, staff, auth, admin, venue.",
        "Service Layer: ChatService xử lý logic chat, gọi AI Agent.",
        "Repository Layer: Truy cập database thông qua SQLAlchemy async.",
        "Schemas: Pydantic models để validate request/response.",
    ]
    for item in backend_items:
        add_bullet_point(doc, item)

    add_body_text(doc, "AI Agent (Domain Layer):", bold=True)

    agent_items = [
        "VenueAgent: Agent chính sử dụng LangChain với Tool Calling.",
        "IntentRouter: Phân loại ý định người dùng (embedding-based).",
        "Tools: query_knowledge, book_court, order_food, call_staff, check_schedule, recommend_menu.",
        "Knowledge Graph Query: Truy vấn Neo4j bằng Cypher, full-text search.",
        "Fallback Mechanism: Tự động chuyển sang LLM dự phòng khi rate limit.",
    ]
    for item in agent_items:
        add_bullet_point(doc, item)

    # 3.3.4. Luồng xử lý yêu cầu
    add_heading_custom(doc, "3.3.4. Luồng xử lý yêu cầu (Request flow)", level=3)

    add_body_text(doc, "Luồng xử lý khi người dùng gửi tin nhắn chat:")

    flow_items = [
        "1. Flutter App gửi POST /api/chat với message và session_id.",
        "2. ChatRouter nhận request, gọi ChatService.process_message().",
        "3. ChatService lấy session history từ Redis.",
        "4. ChatService gọi VenueAgent.process(message, history).",
        "5. IntentRouter kiểm tra intent đơn giản (greeting, goodbye).",
        "6. Nếu phức tạp, gọi LLM AgentExecutor với tools.",
        "7. LLM quyết định gọi tool nào (query_knowledge, book_court, v.v.).",
        "8. Tool thực thi, truy vấn database (Neo4j/PostgreSQL).",
        "9. Kết quả tool trả về LLM để tổng hợp phản hồi.",
        "10. Phản hồi trả về Flutter App, lưu vào Redis session.",
    ]
    for item in flow_items:
        add_body_text(doc, item)

    # 3.3.5. Ưu điểm của kiến trúc
    add_heading_custom(doc, "3.3.5. Ưu điểm của kiến trúc", level=3)

    advantages = [
        "Tách biệt rõ ràng giữa các tầng (Presentation, Application, Domain, Infrastructure).",
        "Dễ dàng mở rộng: thêm tool mới, thêm API endpoint mà không ảnh hưởng các thành phần khác.",
        "Multi-tenant architecture: hỗ trợ nhiều quán thể thao trên cùng hệ thống.",
        "Fallback mechanism: đảm bảo hệ thống hoạt động ngay cả khi LLM chính lỗi.",
        "Sử dụng LLM local (Ollama) giúp giảm chi phí và bảo mật dữ liệu.",
    ]
    for item in advantages:
        add_bullet_point(doc, item)

    # 3.3.6. Hạn chế và hướng phát triển
    add_heading_custom(doc, "3.3.6. Hạn chế và hướng phát triển", level=3)

    limitations = [
        "Hiện tại chỉ hỗ trợ 3 môn thể thao (bida, pickleball, cầu lông), cần mở rộng thêm.",
        "Chưa có hệ thống đánh giá và phản hồi từ khách hàng.",
        "Cần bổ sung analytics và reporting nâng cao cho admin.",
        "Hướng phát triển: tích hợp voice chat, hỗ trợ đa ngôn ngữ, thêm môn thể thao.",
        "Hướng phát triển: triển khai lên cloud (AWS, GCP) với CI/CD pipeline.",
    ]
    for item in limitations:
        add_bullet_point(doc, item)

    # 3.4. Thiết kế API
    add_heading_custom(doc, "3.4. Thiết kế API", level=2)

    # 3.4.1. Xác thực
    add_heading_custom(doc, "3.4.1. Xác thực", level=3)

    add_body_text(doc, "Hệ thống sử dụng JWT (JSON Web Token) để xác thực người dùng.")

    auth_headers = ["Method", "Path", "Mô tả"]
    auth_data = [
        ["POST", "/api/auth/register", "Đăng ký tài khoản mới"],
        ["POST", "/api/auth/login", "Đăng nhập, trả về JWT token"],
    ]
    create_table_with_header(doc, auth_headers, auth_data)

    # 3.4.2. Chat APIs
    add_heading_custom(doc, "3.4.2. Chat APIs", level=3)

    chat_data = [
        ["POST", "/api/chat", "Gửi tin nhắn chat, nhận phản hồi từ AI Agent"],
        ["GET", "/api/chat/history/{session_id}", "Lấy lịch sử chat theo session"],
    ]
    create_table_with_header(doc, auth_headers, chat_data)

    # 3.4.3. Booking APIs
    add_heading_custom(doc, "3.4.3. Booking APIs", level=3)

    booking_api_data = [
        ["POST", "/api/booking/", "Tạo đặt sân mới"],
        ["GET", "/api/booking/{id}", "Lấy thông tin đặt sân"],
        ["GET", "/api/booking/user/{user_id}", "Lấy danh sách đặt sân của user"],
        ["PUT", "/api/booking/{id}/cancel", "Hủy đặt sân"],
        ["GET", "/api/booking/available/", "Kiểm tra sân trống"],
    ]
    create_table_with_header(doc, auth_headers, booking_api_data)

    # 3.4.4. Order APIs
    add_heading_custom(doc, "3.4.4. Order APIs", level=3)

    order_data = [
        ["POST", "/api/order/", "Tạo đơn hàng mới"],
        ["GET", "/api/order/{id}", "Lấy thông tin đơn hàng"],
        ["PUT", "/api/order/{id}/status", "Cập nhật trạng thái đơn hàng"],
    ]
    create_table_with_header(doc, auth_headers, order_data)

    # 3.4.5. Menu APIs
    add_heading_custom(doc, "3.4.5. Menu APIs", level=3)

    menu_api_data = [
        ["GET", "/api/menu/", "Lấy danh sách menu đồ uống (giá VND)"],
    ]
    create_table_with_header(doc, auth_headers, menu_api_data)

    # 3.4.6. Staff APIs
    add_heading_custom(doc, "3.4.6. Staff APIs", level=3)

    staff_data = [
        ["POST", "/api/staff/notify", "Gửi thông báo cho nhân viên"],
        ["POST", "/api/staff-request/", "Tạo yêu cầu hỗ trợ từ khách hàng"],
    ]
    create_table_with_header(doc, auth_headers, staff_data)

    doc.add_paragraph()

    # 3.4.7. Payment APIs
    add_heading_custom(doc, "3.4.7. Payment APIs", level=3)

    add_body_text(doc, "VNPay Payment APIs:", bold=True)

    vnpay_api_data = [
        ["POST", "/api/payment/create", "Tạo URL thanh toán VNPay"],
        ["GET", "/api/payment/callback", "Callback từ VNPay sau thanh toán"],
        ["GET", "/api/payment/query", "Truy vấn trạng thái giao dịch VNPay"],
    ]
    create_table_with_header(doc, auth_headers, vnpay_api_data)

    doc.add_paragraph()
    add_body_text(doc, "Stripe Payment APIs:", bold=True)

    stripe_api_data = [
        ["POST", "/api/stripe/create-checkout", "Tạo Stripe Checkout Session"],
        ["GET", "/api/stripe/success", "Redirect sau thanh toán Stripe thành công"],
        ["GET", "/api/stripe/cancel", "Redirect khi hủy thanh toán Stripe"],
        ["POST", "/api/stripe/webhook", "Webhook xác nhận thanh toán từ Stripe"],
        ["GET", "/api/stripe/config", "Lấy Stripe publishable key"],
    ]
    create_table_with_header(doc, auth_headers, stripe_api_data)

    # 3.4.8. Mô hình bảo mật API
    add_heading_custom(doc, "3.4.8. Mô hình bảo mật API", level=3)

    add_body_text(doc, "Các biện pháp bảo mật được áp dụng:")

    security_items = [
        "JWT Authentication: Mỗi request cần có Bearer token trong header.",
        "Password Hashing: Mật khẩu được mã hóa bằng bcrypt trước khi lưu vào database.",
        "CORS Protection: Chỉ cho phép các origin được whitelist truy cập API.",
        "Input Validation: Tất cả request body được validate bằng Pydantic schemas.",
        "Rate Limiting: Giới hạn số request để tránh abuse (thông qua LLM provider).",
        "Webhook Signature Verification: Xác thực chữ ký webhook từ Stripe và VNPay.",
        "Payment Deduplication: Kiểm tra trùng lặp giao dịch bằng mã giao dịch cổng thanh toán.",
    ]
    for item in security_items:
        add_bullet_point(doc, item)

    add_page_break(doc)

    # ====================================================================
    # CHƯƠNG 4. HIỆN THỰC CHƯƠNG TRÌNH
    # ====================================================================
    add_heading_custom(doc, "CHƯƠNG 4. HIỆN THỰC CHƯƠNG TRÌNH", level=1)

    # 4.1. Cấu trúc thư mục dự án
    add_heading_custom(doc, "4.1. Cấu trúc thư mục dự án", level=2)

    add_body_text(doc,
        "Dự án được tổ chức theo cấu trúc thư mục rõ ràng, tách biệt giữa "
        "backend (Python/FastAPI) và frontend (Flutter/Dart)."
    )

    add_body_text(doc,
        "ChatbotAgent/\n"
        "├── backend/                    # Backend API (Python/FastAPI)\n"
        "│   ├── main.py                 # Entry point\n"
        "│   ├── app/\n"
        "│   │   ├── api/                # API routers\n"
        "│   │   │   ├── auth.py         # Xác thực\n"
        "│   │   │   ├── chat.py         # Chat\n"
        "│   │   │   ├── booking.py      # Đặt sân\n"
        "│   │   │   ├── order.py        # Đơn hàng\n"
        "│   │   │   ├── payment.py      # VNPay payment\n"
        "│   │   │   ├── stripe.py       # Stripe payment\n"
        "│   │   │   └── admin.py        # Quản trị\n"
        "│   │   ├── agent/              # AI Agent core\n"
        "│   │   ├── kg/                 # Knowledge Graph\n"
        "│   │   ├── models/             # SQLAlchemy models\n"
        "│   │   ├── schemas/            # Pydantic schemas\n"
        "│   │   ├── services/           # Business logic\n"
        "│   │   └── core/               # Config, DB, Redis\n"
        "│   ├── data_pipeline/          # KG build pipeline\n"
        "│   └── Dockerfile              # Backend container\n"
        "│\n"
        "├── java-payment-service/       # VNPay Gateway (Java)\n"
        "│   ├── Dockerfile\n"
        "│   ├── pom.xml\n"
        "│   └── src/                    # gRPC + VNPay integration\n"
        "│\n"
        "├── flutter_app/                # Mobile App (Flutter)\n"
        "│   ├── lib/\n"
        "│   │   ├── features/           # Feature modules\n"
        "│   │   │   ├── chat/           # Chat feature\n"
        "│   │   │   ├── booking/        # Booking feature\n"
        "│   │   │   ├── menu/           # Menu feature\n"
        "│   │   │   ├── auth/           # Authentication\n"
        "│   │   │   ├── payment/        # Payment (Stripe/VNPay)\n"
        "│   │   │   ├── admin/          # Admin panel\n"
        "│   │   │   └── profile/        # User profile\n"
        "│   │   ├── core/               # Constants, theme, router\n"
        "│   │   └── shared/             # Shared widgets\n"
        "│   └── pubspec.yaml\n"
        "│\n"
        "├── docker-compose.yml          # Multi-service deployment\n"
        "├── .dockerignore\n"
        "└── proto/                      # gRPC protobuf definitions"
    )

    # 4.2. Hiện thực Backend
    add_heading_custom(doc, "4.2. Hiện thực Backend", level=2)

    # 4.2.1. AI Agent Implementation
    add_heading_custom(doc, "4.2.1. AI Agent Implementation", level=3)

    add_body_text(doc,
        "AI Agent được xây dựng bằng LangChain với kiến trúc Tool Calling. "
        "Agent chính (VenueAgent) sử dụng LLM để phân tích ý định người dùng "
        "và quyết định gọi tool phù hợp."
    )

    add_body_text(doc, "Cấu trúc Agent:")

    agent_structure = [
        "VenueAgent: Agent chính, xử lý chat message.",
        "IntentRouter: Phân loại intent nhanh (embedding-based) cho các câu hỏi đơn giản.",
        "SimpleVenueAgent: Fallback agent khi LLM không khả dụng.",
        "Tools: Các hàm xử lý nghiệp vụ (book_court, order_food, call_staff, query_knowledge).",
    ]
    for item in agent_structure:
        add_bullet_point(doc, item)

    add_body_text(doc, "Luồng xử lý:")

    agent_flow = [
        "1. Nhận message từ người dùng.",
        "2. IntentRouter kiểm tra intent đơn giản (greeting, goodbye).",
        "3. Nếu phức tạp, gọi LLM AgentExecutor.",
        "4. LLM phân tích và gọi tool phù hợp.",
        "5. Tool thực thi, trả kết quả về LLM.",
        "6. LLM tổng hợp phản hồi cuối cùng.",
    ]
    for item in agent_flow:
        add_body_text(doc, item)

    doc.add_paragraph()
    add_body_text(doc, "[Chèn hình Agent Flow Diagram tại đây]", bold=True)

    # ── AI Agent Flow (Mermaid) ──
    add_mermaid_block(doc, "Luồng xử lý AI Agent:", """flowchart TD
    A[Nhận message từ user] --> B{IntentRouter\nphân tích intent}
    B -->|Greeting/Goodbye| C[Trả lời trực tiếp\nkhông gọi LLM]
    B -->|Hỏi luật/kỹ thuật| D[Gọi LLM AgentExecutor]
    B -->|Đặt sân| D
    B -->|Gọi món| D
    B -->|Gọi nhân viên| D

    D --> E[LLM phân tích\nvà chọn tool]

    E -->|query_knowledge| F[Truy vấn Neo4j\nCypher + Fulltext]
    E -->|book_court| G[Kiểm tra availability\nInsert PostgreSQL]
    E -->|order_food| H[Tạo order\nLink booking]
    E -->|call_staff| I[Tạo staff request\nNotify WebSocket]
    E -->|check_schedule| J[Query bookings\ntheo ngày]
    E -->|Không cần tool| K[LLM trả lời trực tiếp]

    F --> L[Kết quả tool\ntrả về LLM]
    G --> L
    H --> L
    I --> L
    J --> L

    L --> M[LLM tổng hợp\nphản hồi cuối cùng]
    K --> M
    C --> M

    M --> N[Trả response\nvề Flutter App]
""")

    # ── Payment Decision Flow (Mermaid) ──
    add_mermaid_block(doc, "Luồng quyết định thanh toán:", """flowchart TD
    A[Đặt sân / Đặt hàng thành công] --> B{Chọn phương thức\nthanh toán?}

    B -->|Stripe| C[POST /stripe/create-payment-intent]
    B -->|VNPay| D[POST /payment/create]
    B -->|Để sau| E[Lưu trạng thái unpaid]

    C --> F[Stripe PaymentIntent\ntrả client_secret]
    F --> G[Flutter: initPaymentSheet]
    G --> H[User nhập thông tin thẻ]
    H --> I{Thanh toán?}
    I -->|Thành công| J[Webhook: payment_intent.succeeded]
    I -->|Thất bại| K[Hiển thị lỗi]
    I -->|Hủy| E

    D --> L[Java Service tạo\nVNPay payment URL]
    L --> M[Flutter: mở VNPay SDK]
    M --> N[User chọn ngân hàng]
    N --> O{Thanh toán?}
    O -->|Thành công| P[Callback: xác thực chữ ký]
    O -->|Thất bại| K
    O -->|Hủy| E

    J --> Q[Cập nhật payment_status = paid]
    P --> Q
    Q --> R[Cập nhật booking/order]
    R --> S[Hiển kết quả thành công]
""")

    # ── User Journey Activity (Mermaid) ──
    add_mermaid_block(doc, "Hoạt động - Journey người dùng (Customer):", """flowchart TD
    START([Mở app]) --> LOGIN{Đăng nhập?}
    LOGIN -->|Chưa có TK| REGISTER[Đăng ký]
    LOGIN -->|Đã có TK| INPUT[Nhập SĐT + MK]
    REGISTER --> INPUT
    INPUT --> HOME[Màn hình chính]

    HOME --> CHAT[Chat AI]
    HOME --> BOOK[Đặt sân]
    HOME --> ORDER[Đặt đồ uống]
    HOME --> STAFF[0 Gọi nhân viên]
    HOME --> PROFILE[Hồ sơ]

    CHAT --> ASK[Nhập tin nhắn]
    ASK --> REPLY[Nhận phản hồi AI]
    REPLY --> CHAT

    BOOK --> SELECT[Chọn loại sân + thời gian]
    SELECT --> CHECK{Sân trống?}
    CHECK -->|Có| CONFIRM[Xác nhận đặt]
    CHECK -->|Không| SELECT
    CONFIRM --> PAY1{Thanh toán?}

    ORDER --> MENU[Xem menu]
    MENU --> CART[Thêm vào giỏ]
    CART --> PLACE[Đặt hàng]
    PLACE --> PAY2{Thanh toán?}

    PAY1 -->|Stripe| STRIPE[Stripe Payment Sheet]
    PAY1 -->|VNPay| VNPAY[VNPay SDK]
    PAY1 -->|Để sau| HOME

    PAY2 -->|Stripe| STRIPE
    PAY2 -->|VNPay| VNPAY
    PAY2 -->|Để sau| HOME

    STRIPE --> RESULT{Kết quả}
    VNPAY --> RESULT
    RESULT -->|OK| SUCCESS[Thành công]
    RESULT -->|Fail| RETRY[Thử lại]
    SUCCESS --> HOME
    RETRY --> HOME

    STAFF --> REQ[Chọn loại yêu cầu]
    REQ --> SEND[Gửi yêu cầu]
    SEND --> WAIT[Chờ nhân viên xử lý]
    WAIT --> HOME
""")

    # 4.2.2. Knowledge Graph Pipeline
    add_heading_custom(doc, "4.2.2. Knowledge Graph Pipeline", level=3)

    add_body_text(doc,
        "Data Pipeline chuyển đổi dữ liệu thô (text files) thành Knowledge Graph "
        "trong Neo4j qua 4 bước:"
    )

    pipeline_steps = [
        "Step 1 - Scrape/Collect: Thu thập dữ liệu từ web (BeautifulSoup), PDF (pdfplumber), YouTube (yt-dlp).",
        "Step 2 - Extract Entities: Sử dụng LLM (Ollama qwen2.5-coder:7b) để trích xuất entities và relationships.",
        "Step 3 - Build Graph: Loại bỏ trùng lặp, tạo constraints/indexes, insert vào Neo4j.",
        "Step 4 - Embed Nodes: Tạo vector embeddings cho entity names/descriptions.",
    ]
    for step in pipeline_steps:
        add_bullet_point(doc, step)

    doc.add_paragraph()
    add_body_text(doc, "Kết quả extraction:", bold=True)

    extraction_headers = ["Metric", "Value"]
    extraction_data = [
        ["Files processed", "10"],
        ["Total chunks", "83"],
        ["Entities extracted", "563"],
        ["Relationships extracted", "464"],
        ["Unique entities (after dedup)", "418"],
        ["Unique relationships (after dedup)", "441"],
        ["Processing time", "~91 minutes"],
    ]
    create_table_with_header(doc, extraction_headers, extraction_data)

    # 4.2.3. Database Models
    add_heading_custom(doc, "4.2.3. Database Models", level=3)

    add_body_text(doc, "Các model chính trong hệ thống (SQLAlchemy async):")

    model_items = [
        "User: Quản lý tài khoản người dùng (phone, name, email, role, password_hash, stripe_customer_id).",
        "Booking: Quản lý đặt sân (user_id, venue_id, court_type, start_time, end_time, status, total_price, payment_status).",
        "Order: Quản lý đơn hàng (user_id, venue_id, total_price, status, payment_status).",
        "Venue: Quản lý quán thể thao (name, address, phone, business_id).",
        "MenuItem: Quản lý menu đồ uống (name, price, category, available).",
        "StaffRequest: Quản lý yêu cầu hỗ trợ từ khách hàng.",
        "Payment: Quản lý giao dịch thanh toán (order_type, order_id, amount, bank_code, status, vnp_transaction_no).",
    ]
    for item in model_items:
        add_bullet_point(doc, item)

    # 4.2.4. API Endpoints
    add_heading_custom(doc, "4.2.4. API Endpoints", level=3)

    add_body_text(doc, "Hệ thống cung cấp các nhóm API chính:")

    api_groups = [
        "Authentication: POST /api/auth/register, POST /api/auth/login",
        "Chat: POST /api/chat, GET /api/chat/history/{session_id}",
        "Booking: POST /api/booking/, GET /api/booking/{id}, PUT /api/booking/{id}/cancel",
        "Order: POST /api/order/, GET /api/order/{id}, PUT /api/order/{id}/status",
        "Menu: GET /api/menu/",
        "Staff: POST /api/staff/notify, POST /api/staff-request/",
        "VNPay Payment: POST /api/payment/create, GET /api/payment/callback, GET /api/payment/query",
        "Stripe Payment: POST /api/stripe/create-checkout, POST /api/stripe/webhook, GET /api/stripe/config",
        "Venue: CRUD quản lý quán thể thao",
        "Admin: Quản lý hệ thống (chỉ admin), quản lý giá dịch vụ",
    ]
    for item in api_groups:
        add_bullet_point(doc, item)

    doc.add_paragraph()

    # 4.2.5. Payment Implementation
    add_heading_custom(doc, "4.2.5. Payment Implementation", level=3)

    add_body_text(doc,
        "Hệ thống thanh toán được thiết kế theo kiến trúc microservice, "
        "hỗ trợ hai cổng thanh toán: Stripe (quốc tế) và VNPay (trong nước)."
    )

    add_body_text(doc, "Kiến trúc thanh toán:", bold=True)

    payment_arch = [
        "PaymentService: Service chính xử lý logic thanh toán, quản lý trạng thái payment.",
        "Stripe Integration: Backend tạo Checkout Session bằng Stripe Python SDK, frontend mở WebView.",
        "VNPay Integration: Java microservice đóng vai trò proxy VNPay, giao tiếp với backend qua gRPC.",
        "Webhook Handling: Stripe webhook xác nhận server-to-server, VNPay callback redirect người dùng.",
        "Payment Status Flow: unpaid -> pending -> paid/failed, cập nhật lên booking/order.",
    ]
    for item in payment_arch:
        add_bullet_point(doc, item)

    doc.add_paragraph()
    add_body_text(doc, "Luồng thanh toán Stripe:", bold=True)

    stripe_flow = [
        "1. Người dùng chọn 'Thanh toán ngay' sau khi đặt sân/đồ uống.",
        "2. Frontend gọi POST /api/stripe/create-checkout với order_id, amount.",
        "3. Backend tạo Stripe Checkout Session, trả về checkout_url.",
        "4. Frontend mở WebView hiển thị trang thanh toán Stripe.",
        "5. Người dùng nhập thông tin thẻ và hoàn tất thanh toán.",
        "6. Stripe gửi webhook POST /api/stripe/webhook xác nhận thanh toán.",
        "7. Backend xác thực chữ ký webhook, cập nhật payment_status = 'paid'.",
        "8. Frontend detect URL redirect, hiển thị kết quả thanh toán.",
    ]
    for item in stripe_flow:
        add_body_text(doc, item)

    doc.add_paragraph()
    add_body_text(doc, "Luồng thanh toán VNPay:", bold=True)

    vnpay_flow = [
        "1. Người dùng chọn thanh toán qua VNPay.",
        "2. Frontend gọi POST /api/payment/create với order_id, amount.",
        "3. Backend gọi Java Payment Service tạo VNPay payment URL.",
        "4. Frontend mở VNPay Native SDK (Method Channel) với payment URL.",
        "5. Người dùng chọn ngân hàng và hoàn tất thanh toán.",
        "6. VNPay redirect về callback URL, backend xác thực chữ ký.",
        "7. Backend cập nhật payment_status = 'paid', trả kết quả.",
    ]
    for item in vnpay_flow:
        add_body_text(doc, item)

    # 4.2.6. Docker Deployment
    add_heading_custom(doc, "4.2.6. Docker Deployment", level=3)

    add_body_text(doc,
        "Hệ thống được đóng gói thành 4 Docker services, quản lý bởi Docker Compose:"
    )

    docker_headers = ["Service", "Image", "Port", "Mô tả"]
    docker_data = [
        ["postgres", "postgres:16-alpine", "internal", "PostgreSQL database, healthcheck"],
        ["redis", "redis:7-alpine", "internal", "Cache & message broker, healthcheck"],
        ["java-payment", "custom build", "internal:9090", "VNPay gateway, giao tiếp gRPC"],
        ["python-backend", "custom build", "8000:8000", "FastAPI API server"],
    ]
    create_table_with_header(doc, docker_headers, docker_data)

    doc.add_paragraph()
    add_body_text(doc, "Đặc điểm triển khai Docker:", bold=True)

    docker_features = [
        "Mạng nội bộ (bridge network): Các service giao tiếp qua mạng 'internal', chỉ backend expose ra ngoài.",
        "Health checks: PostgreSQL và Redis có health check, backend chỉ khởi động khi dependencies sẵn sàng.",
        "Volume persistence: Dữ liệu PostgreSQL được persist qua Docker volume 'pgdata'.",
        "Environment variables: Cấu hình qua .env file, bao gồm database URL, Redis URL, Stripe keys, VNPay config.",
    ]
    for item in docker_features:
        add_bullet_point(doc, item)

    # 4.3. Hiện thực Frontend
    add_heading_custom(doc, "4.3. Hiện thực Frontend (Flutter App)", level=2)

    # 4.3.1. Cấu trúc Feature-based
    add_heading_custom(doc, "4.3.1. Cấu trúc Feature-based", level=3)

    add_body_text(doc,
        "Ứng dụng Flutter được tổ chức theo kiến trúc feature-based, "
        "mỗi feature là một module độc lập với 3 tầng:"
    )

    feature_layers = [
        "data/: Repository, data sources, models (DTOs).",
        "domain/: Business logic, entities, use cases.",
        "presentation/: UI widgets, state management (Riverpod).",
    ]
    for item in feature_layers:
        add_bullet_point(doc, item)

    add_body_text(doc, "Các feature chính:")

    features = [
        "Chat Feature: Giao diện chat với AI, hỗ trợ markdown rendering.",
        "Booking Feature: Chọn loại sân, thời gian, xem sân trống, xác nhận đặt, thanh toán.",
        "Menu Feature: Hiển thị danh sách đồ uống, đặt hàng.",
        "Auth Feature: Đăng nhập/đăng ký với số điện thoại.",
        "Payment Feature: Thanh toán qua Stripe (Native Payment Sheet) và VNPay (Native SDK).",
        "Billing Feature: Xem hóa đơn đặt sân tổng hợp (phí sân + đơn hàng).",
        "Profile Feature: Quản lý thông tin cá nhân, đổi mật khẩu, đăng xuất.",
        "Admin Feature: Quản lý booking, billing, resource pricing, dashboard analytics.",
        "Staff Request Feature: Gửi và quản lý yêu cầu hỗ trợ nhân viên.",
        "Staff Chat Feature: Chat thời gian thực giữa nhân viên và khách hàng.",
        "Notification Feature: Thông báo real-time qua WebSocket cho tất cả role.",
    ]
    for item in features:
        add_bullet_point(doc, item)

    # 4.3.2. State Management
    add_heading_custom(doc, "4.3.2. State Management với Riverpod", level=3)

    add_body_text(doc,
        "Ứng dụng sử dụng Riverpod để quản lý trạng thái. Riverpod cung cấp "
        "các ưu điểm: type-safe, testable, không phụ thuộc BuildContext. "
        "Các provider chính: AuthProvider, ChatProvider, BookingProvider, MenuProvider."
    )

    # 4.3.3. Navigation
    add_heading_custom(doc, "4.3.3. Navigation với GoRouter", level=3)

    add_body_text(doc,
        "Điều hướng trong ứng dụng sử dụng GoRouter, hỗ trợ declarative routing "
        "và deep linking. Các route chính: /login, /home, /chat, /booking, /menu, /profile."
    )

    # 4.3.4. Chat UI
    add_heading_custom(doc, "4.3.4. Chat UI với flutter_chat_ui", level=3)

    add_body_text(doc,
        "Giao diện chat sử dụng thư viện flutter_chat_ui, hỗ trợ: "
        "hiển thị tin nhắn văn bản, markdown rendering (flutter_markdown), "
        "typing indicator, message bubbles, và scroll-to-bottom."
    )

    # 4.3.5. Giới thiệu giao diện ứng dụng
    add_heading_custom(doc, "4.3.5. Giới thiệu giao diện ứng dụng", level=3)

    add_body_text(doc,
        "Phần này giới thiệu các màn hình chính của ứng dụng Sports Venue Chatbot. "
        "Mỗi màn hình được mô tả chức năng và cách sử dụng. "
        "Ảnh chụp màn hình tương ứng được đính kèm bên dưới mỗi mục."
    )

    doc.add_paragraph()
    add_body_text(doc, "A. Màn hình xác thực (Authentication)", bold=True)

    auth_screens = [
        ("Hình 4.7 - Màn hình đăng nhập:", "Người dùng nhập số điện thoại và mật khẩu để đăng nhập. "
         "Hỗ trợ chuyển sang màn hình đăng ký cho người dùng mới. "
         "Giao diện đơn giản, tập trung vào trải nghiệm người dùng."),
    ]
    for title, desc in auth_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "B. Màn hình chính và điều hướng", bold=True)

    home_screens = [
        ("Hình 4.8 - Màn hình chính (Home Screen):", "Màn hình chính sau khi đăng nhập, hiển thị "
         "các chức năng chính: Chat AI, Đặt sân, Menu, Gọi nhân viên. "
         "Thanh điều hướng dưới cùng cho phép chuyển đổi giữa các tab."),
    ]
    for title, desc in home_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "C. Chat với AI Agent", bold=True)

    chat_screens = [
        ("Hình 4.9 - Giao diện Chat với AI Agent:", "Màn hình chat chính, người dùng nhập tin nhắn "
         "và nhận phản hồi từ AI. Hỗ trợ markdown rendering cho định dạng văn bản, "
         "hiển thị typing indicator khi AI đang xử lý. "
         "Có thể hỏi về luật chơi, kỹ thuật, đặt sân, gọi đồ uống trực tiếp qua chat."),
    ]
    for title, desc in chat_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "D. Đặt sân (Booking)", bold=True)

    booking_screens = [
        ("Hình 4.10 - Màn hình đặt sân (Booking):", "Chọn loại sân (bida, pickleball, cầu lông), "
         "xem danh sách sân trống theo thời gian. Hiển thị giá sân và trạng thái availability."),
        ("Hình 4.11 - Màn hình chọn thời gian đặt sân:", "Chọn ngày và giờ bắt đầu, kết thúc. "
         "Hệ thống kiểm tra tự động sân trống và hiển thị kết quả real-time."),
    ]
    for title, desc in booking_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "E. Menu và đặt đồ uống", bold=True)

    menu_screens = [
        ("Hình 4.12 - Màn hình Menu đồ uống:", "Hiển thị danh sách đồ uống/đồ ăn theo danh mục, "
         "có hình ảnh, giá tiền (VND). Hỗ trợ tìm kiếm và lọc theo loại."),
        ("Hình 4.13 - Màn hình giỏ hàng (Cart):", "Xem các món đã chọn, điều chỉnh số lượng, "
         "xem tổng tiền. Nút 'Đặt hàng' để xác nhận đơn."),
    ]
    for title, desc in menu_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "F. Thanh toán", bold=True)

    payment_screens = [
        ("Hình 4.14 - Dialog chọn phương thức thanh toán:", "Sau khi đặt sân hoặc đặt đồ uống, "
         "hệ thống hiển thị dialog cho phép chọn phương thức: Stripe (quốc tế) hoặc VNPay (trong nước). "
         "Có nút 'Để sau' nếu người dùng chưa muốn thanh toán ngay."),
        ("Hình 4.15 - Màn hình thanh toán Stripe (Payment Sheet):", "Stripe Native Payment Sheet "
         "hiển thị trực tiếp trong app (không qua WebView). Người dùng nhập thông tin thẻ "
         "và xác nhận thanh toán. Hỗ trợ Visa, Mastercard, JCB."),
        ("Hình 4.16 - Màn hình thanh toán VNPay:", "VNPay Native SDK mở trang thanh toán "
         "trên thiết bị. Người dùng chọn ngân hàng và hoàn tất giao dịch."),
        ("Hình 4.17 - Kết quả thanh toán thành công:", "Hiển thị kết quả thanh toán "
         "với mã giao dịch, số tiền, và trạng thái. Có nút quay về trang chính."),
    ]
    for title, desc in payment_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "G. Gọi nhân viên và quản lý yêu cầu", bold=True)

    staff_screens = [
        ("Hình 4.18 - Màn hình gọi nhân viên (Call Staff Dialog):", "Dialog cho phép khách hàng "
         "chọn loại yêu cầu (bảo trì, dọn dẹp, gọi đồ uống, hỗ trợ khác) "
         "và nhập mô tả chi tiết. Gửi yêu cầu đến nhân viên phụ trách."),
        ("Hình 4.19 - Màn hình danh sách yêu cầu (Staff Requests):", "Danh sách các yêu cầu "
         "đã gửi, hiển thị trạng thái: chờ xử lý, đã chấp nhận, đã hoàn thành."),
        ("Hình 4.20 - Màn hình quản lý yêu cầu (Staff Request Management):", "Giao diện cho nhân viên "
         "xem và xử lý yêu cầu từ khách hàng. Chia thành 2 nhóm: 'Chờ xử lý' và 'Đang xử lý'. "
         "Nhân viên có thể chấp nhận, hoàn thành, hoặc hủy yêu cầu."),
    ]
    for title, desc in staff_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "H. Chat với nhân viên (Staff Chat)", bold=True)

    staff_chat_screens = [
        ("Hình 4.21 - Màn hình chat với nhân viên (Staff Chat):", "Giao diện chat thời gian thực "
         "giữa khách hàng và nhân viên. Sử dụng WebSocket để truyền tin nhắn tức thì. "
         "Hỗ trợ hiển thị tin nhắn, thời gian gửi, và trạng thái đối phương online."),
        ("Hình 4.22 - Màn hình hộp thư nhân viên (Staff Inbox):", "Danh sách các cuộc trò chuyện "
         "của nhân viên với khách hàng. Hiển thị tin nhắn cuối cùng, thời gian, "
         "và số tin nhắn chưa đọc."),
    ]
    for title, desc in staff_chat_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "I. Thông báo vận hành", bold=True)

    notif_screens = [
        ("Hình 4.23 - Màn hình thông báo vận hành (Notifications):", "Danh sách thông báo "
         "cho nhân viên/admin: đặt sân mới, đặt đồ uống mới, yêu cầu hỗ trợ. "
         "Hiển thị kết nối realtime, đánh dấu đã đọc, phân trang vô hạn."),
    ]
    for title, desc in notif_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "J. Hóa đơn và lịch sử", bold=True)

    billing_screens = [
        ("Hình 4.24 - Màn hình hóa đơn đặt sân (Booking Bill):", "Hiển thị hóa đơn tổng hợp "
         "cho một lần đặt sân: phí sân theo giờ, danh sách đồ uống đã đặt, "
         "và tổng tiền thanh toán. Hỗ trợ xem chi tiết từng đơn hàng."),
        ("Hình 4.25 - Màn hình lịch sử đơn hàng (Order History):", "Danh sách tất cả đơn hàng "
         "của khách hàng, hiển thị trạng thái, tổng tiền, và nút thanh toán nếu chưa trả."),
    ]
    for title, desc in billing_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "K. Quản trị (Admin)", bold=True)

    admin_screens = [
        ("Hình 4.26 - Màn hình quản lý đặt sân (Admin Booking):", "Admin xem tất cả đặt sân, "
         "lọc theo trạng thái, xem hóa đơn chi tiết, cập nhật trạng thái."),
        ("Hình 4.27 - Màn hình quản lý menu (Admin Menu):", "Admin thêm/sửa/xóa món trong menu, "
         "cập nhật giá, bật/tắt tình trạng còn phục vụ."),
        ("Hình 4.28 - Màn hình Dashboard Admin:", "Tổng quan doanh thu, số lượng đặt sân, "
         "đơn hàng theo ngày. Biểu đồ và thống kê."),
    ]
    for title, desc in admin_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    add_body_text(doc, "L. Hồ sơ cá nhân và điều hướng nhân viên", bold=True)

    profile_screens = [
        ("Hình 4.29 - Màn hình hồ sơ cá nhân (Profile):", "Xem thông tin cá nhân, "
         "đổi mật khẩu, đăng xuất. Hiển thị tên, số điện thoại, email."),
        ("Hình 4.30 - Thanh điều hướng nhân viên (Staff Shell):", "Thanh navigation "
         "dành cho nhân viên với 3 tab: Thông báo, Tin nhắn (Inbox), Yêu cầu. "
         "Hiển thị badge số thông báo chưa đọc."),
    ]
    for title, desc in profile_screens:
        add_body_text(doc, title, bold=True)
        add_body_text(doc, desc)
        add_body_text(doc, "[Chèn ảnh chụp màn hình tại đây]", bold=True)
        doc.add_paragraph()

    # 4.4. Môi trường triển khai
    add_heading_custom(doc, "4.4. Môi trường triển khai", level=2)

    deploy_headers = ["Thành phần", "Mô tả"]
    deploy_data = [
        ["Backend Server", "FastAPI chạy trên localhost:8000 (development), Docker container"],
        ["Ollama", "LLM local trên localhost:11434, model qwen2.5-coder:7b"],
        ["Neo4j AuraDB", "Cloud-hosted Knowledge Graph (free tier)"],
        ["PostgreSQL", "Docker container (postgres:16-alpine) hoặc WSL (Ubuntu)"],
        ["Redis", "Docker container (redis:7-alpine) hoặc WSL (Ubuntu)"],
        ["Java Payment Service", "Docker container, port 9090, VNPay gateway"],
        ["Stripe", "Stripe API (test mode), webhook endpoint"],
        ["Flutter App", "Chạy trên Android emulator hoặc thiết bị thật"],
        ["Docker Compose", "Quản lý 4 services: postgres, redis, java-payment, python-backend"],
    ]
    create_table_with_header(doc, deploy_headers, deploy_data)

    # 4.5. Kết quả đạt được
    add_heading_custom(doc, "4.5. Kết quả đạt được", level=2)

    results = [
        "Hệ thống chatbot AI hoạt động, trả lời được câu hỏi về luật chơi và kỹ thuật bida, pickleball, cầu lông.",
        "Knowledge Graph với 418 entities và 441 relationships được xây dựng thành công trên Neo4j.",
        "Chức năng đặt sân hoạt động, kiểm tra sân trống theo thời gian thực.",
        "Chức năng đặt đồ uống hoạt động, hiển thị menu với giá VND.",
        "Chức năng gọi nhân viên hoạt động qua hệ thống thông báo real-time.",
        "Tích hợp thanh toán trực tuyến: Stripe (quốc tế) và VNPay (trong nước) hoạt động.",
        "Hệ thống xác nhận thanh toán qua webhook (Stripe) và callback (VNPay).",
        "Admin quản lý giá dịch vụ, xem hóa đơn chi tiết booking với tổng tiền.",
        "Ứng dụng Flutter chạy được trên cả Android và iOS.",
        "Hệ thống có fallback mechanism khi LLM chính lỗi.",
        "Triển khai bằng Docker Compose với 4 dịch vụ containerized.",
    ]
    for item in results:
        add_bullet_point(doc, item)

    # 4.6. Hướng phát triển
    add_heading_custom(doc, "4.6. Hướng phát triển", level=2)

    future = [
        "Hỗ trợ voice chat (speech-to-text, text-to-speech).",
        "Thêm nhiều môn thể thao hơn (tennis, bóng bàn, v.v.).",
        "Triển khai lên cloud (AWS, GCP) với CI/CD pipeline.",
        "Xây dựng hệ thống analytics và reporting nâng cao cho admin.",
        "Hỗ trợ đa ngôn ngữ (Vietnamese, English).",
        "Tích hợp hệ thống đánh giá và phản hồi từ khách hàng.",
        "Phát triển ứng dụng web version (Flutter Web).",
        "Tích hợp thêm cổng thanh toán MoMo.",
    ]
    for item in future:
        add_bullet_point(doc, item)

    add_page_break(doc)

    # ====================================================================
    # TÀI LIỆU THAM KHẢO
    # ====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TÀI LIỆU THAM KHẢO")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    references = [
        "[1] FastAPI Documentation. https://fastapi.tiangolo.com/",
        "[2] Flutter Documentation. https://docs.flutter.dev/",
        "[3] LangChain Documentation. https://python.langchain.com/",
        "[4] Neo4j Documentation. https://neo4j.com/docs/",
        "[5] Ollama Documentation. https://ollama.ai/",
        "[6] PostgreSQL Documentation. https://www.postgresql.org/docs/",
        "[7] Redis Documentation. https://redis.io/documentation/",
        "[8] Riverpod Documentation. https://riverpod.dev/",
        "[9] GoRouter Documentation. https://pub.dev/packages/go_router",
        "[10] WPA Rules. World Pool-Billiard Association. https://wpapool.com/",
        "[11] USAPA Rules. USA Pickleball Association. https://usapickleball.org/",
        "[12] BWF Laws. Badminton World Federation. https://bwfbadminton.com/",
        "[13] Stripe Documentation. https://stripe.com/docs",
        "[14] VNPay Documentation. https://sandbox.vnpayment.vn/apis/",
        "[15] Docker Documentation. https://docs.docker.com/",
        "[16] Docker Compose Documentation. https://docs.docker.com/compose/",
    ]

    for ref in references:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        set_paragraph_format(p, line_spacing=1.5)

    add_page_break(doc)

    # ====================================================================
    # PHỤ LỤC
    # ====================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PHỤ LỤC")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    doc.add_paragraph()

    # Phụ lục A
    add_body_text(doc, "Phụ lục A: Source code", bold=True)

    doc.add_paragraph()
    add_body_text(doc, "A.1. Backend - main.py", bold=True)
    add_body_text(doc,
        "[Chèn source code main.py tại đây hoặc tham chiếu đến repository]"
    )

    doc.add_paragraph()
    add_body_text(doc, "A.2. Backend - AI Agent", bold=True)
    add_body_text(doc,
        "[Chèn source code agent.py tại đây]"
    )

    doc.add_paragraph()
    add_body_text(doc, "A.3. Frontend - Chat Screen", bold=True)
    add_body_text(doc,
        "[Chèn source code chat_screen.dart tại đây]"
    )

    doc.add_paragraph()

    # Phụ lục B
    add_body_text(doc, "Phụ lục B: API Documentation", bold=True)
    add_body_text(doc,
        "[Chèn Swagger UI screenshot hoặc API documentation tại đây]"
    )

    doc.add_paragraph()

    # Phụ lục C
    add_body_text(doc, "Phụ lục C: Ảnh chụp màn hình ứng dụng", bold=True)
    add_body_text(doc,
        "[Chèn ảnh chụp màn hình các màn hình chính của ứng dụng tại đây]"
    )

    # ===== SAVE =====
    output_path = os.path.join(
        os.path.dirname(__file__),
        "BaoCao_DoAnTotNghiep_ChatbotAI.docx"
    )
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_report()
